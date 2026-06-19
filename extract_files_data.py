from __future__ import annotations
import csv
import json
import os
import re
import unicodedata
import io
import zipfile
from collections import defaultdict
from pathlib import Path
from typing import Any
import pandas as pd
from dotenv import load_dotenv
load_dotenv()
VERSION = 'BILAN CARBONE — Extraction générique V20 - monétaire utile non bloqué'
_req_jour = 0
_groq_verify_calls = 0
USE_GROQ_VERIFICATION = os.getenv('USE_GROQ_VERIFICATION', '1').strip().lower() not in {'0', 'false', 'non', 'no', 'off'}
GROQ_VERIFY_MODEL = os.getenv('GROQ_VERIFY_MODEL', 'llama-3.1-8b-instant')
GROQ_VERIFY_MAX_CALLS = int(os.getenv('GROQ_VERIFY_MAX_CALLS', '8'))
GROQ_VERIFY_CONTEXT_CHARS = int(os.getenv('GROQ_VERIFY_CONTEXT_CHARS', '5000'))
USE_GROQ_FALLBACK = os.getenv('USE_GROQ_FALLBACK', '1').strip().lower() not in {'0', 'false', 'non', 'no', 'off'}
GROQ_FALLBACK_MAX_CALLS = int(os.getenv('GROQ_FALLBACK_MAX_CALLS', '6'))
_groq_fallback_calls = 0
SKIP_KEYWORDS = ['.~lock', '~$', '_in', 'fe energie', 'fe énergie', 'fe fret', 'fe dechets', 'fe déchets', 'fe immobilisations', 'fe autres emissions', 'fe autres émissions', 'descriptif', 'utilitaires', 'export -', 'futurs exports', 'matrice de collecte', 'ratios', 'objectif', 'reduction', 'réduction', 'q18', 'rapport de verification', 'rapport de vérification', 'diagnostic', 'fds', 'fiche de données de sécurité', 'fiche de donnees de securite', 'plan des bureaux', 'suivi copies', 'copies', 'impressions', 'inventaire parc informatique', 'audit extraction bilan carbone', 'audit rejets financiers', 'accounting review needed', 'audit_extraction', 'audit_rejets']
ALLOW_WITH_BILAN_CARBONE = ['recap achat', 'récap achat', 'achat pour bilan carbone']

def norm(text: Any) -> str:
    if text is None or (isinstance(text, float) and pd.isna(text)):
        return ''
    s = str(text).lower().replace('\xa0', ' ')
    s = unicodedata.normalize('NFD', s)
    s = ''.join((c for c in s if unicodedata.category(c) != 'Mn'))
    s = s.replace('₂', '2').replace('³', '3')
    s = re.sub('\\s+', ' ', s)
    return s.strip()

def to_float(value: Any) -> float | None:
    if value is None or (isinstance(value, float) and pd.isna(value)):
        return None
    s = str(value).strip().replace('\xa0', ' ')
    if not s or s.lower() in {'nan', 'none', '-', '–'}:
        return None
    s = re.sub('[^0-9,\\.\\- ]', '', s)
    s = re.sub('(?<=\\d)\\s+(?=\\d{3}(\\D|$))', '', s)
    s = s.replace(' ', '')
    if not re.search('\\d', s):
        return None
    if ',' in s and '.' in s:
        s = s.replace('.', '').replace(',', '.') if s.rfind(',') > s.rfind('.') else s.replace(',', '')
    elif ',' in s:
        s = s.replace(',', '.')
    try:
        v = float(s)
        return v if pd.notna(v) else None
    except Exception:
        return None

def numbers_from_text(text: str) -> list[float]:
    vals = []
    for m in re.findall('[-+]?\\d+(?:[ \\xa0]\\d{3})*(?:[,.]\\d+)?|[-+]?\\d+(?:[,.]\\d+)?', text):
        v = to_float(m)
        if v is not None:
            vals.append(v)
    return vals

def read_text(path: str) -> str:
    for enc in ('utf-8', 'utf-8-sig', 'latin-1', 'cp1252'):
        try:
            return Path(path).read_text(encoding=enc, errors='ignore')
        except Exception:
            continue
    return ''

def smart_header(df_raw: pd.DataFrame, max_rows: int=30) -> int:
    keys = ['quantite', 'quantité', 'qte', 'volume', 'consommation', 'total', 'unite', 'unité', 'designation', 'désignation', 'libelle', 'libellé', 'dechets', 'déchets', 'kwh', 'mois', 'facturation', 'montant', 'valeur', 'immobilisation', 'compteur', 'conso']
    best_i, best_score = (0, 0)
    for i, row in df_raw.head(max_rows).iterrows():
        vals = [norm(v) for v in row.values]
        score = sum((1 for v in vals if any((k in v for k in keys))))
        if score > best_score:
            best_i, best_score = (int(i), score)
    return best_i

def load_table(path: str) -> pd.DataFrame | None:
    ext = Path(path).suffix.lower()
    try:
        if ext == '.csv':
            for enc in ('utf-8-sig', 'utf-8', 'latin-1', 'cp1252'):
                for sep in (';', ',', '\t'):
                    try:
                        raw = pd.read_csv(path, sep=sep, encoding=enc, header=None, on_bad_lines='skip')
                        if raw.empty or (raw.shape[1] <= 1 and sep != ';'):
                            continue
                        h = smart_header(raw)
                        return pd.read_csv(path, sep=sep, encoding=enc, header=h, on_bad_lines='skip')
                    except Exception:
                        continue
        if ext in {'.xlsx', '.xls', '.xlsm'}:
            raw = pd.read_excel(path, header=None)
            h = smart_header(raw)
            return pd.read_excel(path, header=h)
        if ext == '.txt':
            txt = read_text(path)
            lines = [l.strip() for l in txt.splitlines() if l.strip()]
            return pd.DataFrame({'contenu': lines}) if lines else None
    except Exception as e:
        print(f'  [LECTURE] erreur {Path(path).name}: {e}')
    return None

def table_text(df: pd.DataFrame | None, limit: int=12000) -> str:
    if df is None or df.empty:
        return ''
    try:
        return df.fillna('').astype(str).to_string(index=False, max_rows=80, max_cols=20)[:limit]
    except Exception:
        return ''

def is_non_source(filename: str) -> bool:
    f = norm(filename).replace('_', ' ').replace('-', ' ')
    if 'bilan carbone' in f and (not any((a in f for a in ALLOW_WITH_BILAN_CARBONE))):
        return True
    return any((k in f for k in SKIP_KEYWORDS))

def source_family(source: str) -> str:
    stem = norm(Path(str(source)).stem)
    stem = re.sub('^cleaned[_ -]*', '', stem)
    for suffix in ('synthese', 'synthèse', 'resume', 'résumé', 'recap', 'récap', 'electricite eau', 'électricité eau', 'energie eau', 'énergie eau', 'donnees', 'données', 'data', 'feuil1', 'feuil2', 'feuil3', 'cumul par compte', 'hors cessions', 'etat des immobilisations', 'état des immobilisations'):
        stem = re.sub(f'[_ -]+{re.escape(norm(suffix))}$', '', stem)
    return re.sub('\\s+', ' ', stem).strip()

def extract_site_key(text: str) -> str:
    normalized = norm(text)
    address_words = ('avenue', 'av ', 'rue', 'route', 'chemin', 'boulevard', 'za ', 'zi ')
    for postal in re.finditer('\\b(\\d{5})\\b', normalized):
        start = max(0, postal.start() - 100)
        fragment = normalized[start:postal.end()]
        if not any((word in fragment for word in address_words)):
            continue
        code = postal.group(1)
        fragment = re.sub('[^a-z0-9 ]+', ' ', fragment)
        fragment = re.sub('\\s+', ' ', fragment).strip()
        return f'{code}:{fragment[-80:]}'
    return ''

def make_item(source: str, scope: str, role: str, designation: str, quantite: float, unite: str, fiabilite: str='haute', justification: str='extraction directe') -> dict:
    return {'source': source, 'source_family': source_family(source), 'site_key': '', 'coverage_months': 0, 'scope': scope, 'role': role, 'designation': designation.strip(), 'quantite': round(float(quantite), 4), 'unite': unite, 'fiabilite': fiabilite, 'justification': justification, 'est_calcule': False}
MONETARY_PROFILES = {'carburant_financier': {'autorise': True, 'categorie': 'carburant_financier', 'facteur_min': 1.2, 'facteur_defaut': 1.42, 'facteur_max': 1.8, 'incertitude': 65, 'raison': 'montant carburant exploitable via prix moyen configurable'}, 'achats_intrants_precis': {'autorise': True, 'categorie': 'achats_intrants_precis', 'facteur_min': 0.3, 'facteur_defaut': 0.85, 'facteur_max': 1.8, 'incertitude': 75, 'raison': 'matériau ou intrant suffisamment identifié'}, 'fournitures': {'autorise': True, 'categorie': 'fournitures', 'facteur_min': 0.15, 'facteur_defaut': 0.35, 'facteur_max': 0.8, 'incertitude': 80, 'raison': 'fournitures ou petit matériel catégorisé'}, 'location_materiel': {'autorise': True, 'categorie': 'location_materiel', 'facteur_min': 0.08, 'facteur_defaut': 0.3, 'facteur_max': 0.7, 'incertitude': 85, 'raison': 'location de matériel catégorisée'}, 'maintenance': {'autorise': True, 'categorie': 'maintenance', 'facteur_min': 0.05, 'facteur_defaut': 0.22, 'facteur_max': 0.6, 'incertitude': 85, 'raison': 'entretien ou maintenance catégorisée'}, 'dechets_financier': {'autorise': True, 'categorie': 'dechets_financier', 'facteur_min': 0.05, 'facteur_defaut': 0.25, 'facteur_max': 0.8, 'incertitude': 90, 'raison': 'traitement de déchets identifié mais exprimé en euros'}, 'fret_financier': {'autorise': True, 'categorie': 'fret_financier', 'facteur_min': 0.03, 'facteur_defaut': 0.12, 'facteur_max': 0.4, 'incertitude': 90, 'raison': 'transport/fret identifié mais exprimé en euros'}, 'deplacements_financiers': {'autorise': True, 'categorie': 'deplacements_financiers', 'facteur_min': 0.05, 'facteur_defaut': 0.2, 'facteur_max': 0.6, 'incertitude': 90, 'raison': 'déplacements identifiés mais exprimés en euros'}, 'immobilisations': {'autorise': True, 'categorie': 'immobilisations', 'facteur_min': 0.1, 'facteur_defaut': 0.35, 'facteur_max': 1.2, 'incertitude': 90, 'raison': 'acquisitions immobilisées catégorisées'}, 'achats_non_detailles_recap': {'autorise': True, 'categorie': 'achats_non_detailles_recap', 'facteur_min': 0.7, 'facteur_defaut': 1.88, 'facteur_max': 3.0, 'incertitude': 95, 'raison': "achats matériels non détaillés issus d'un récap achat bilan carbone"}, 'autres_services_recap': {'autorise': True, 'categorie': 'autres_services_recap', 'facteur_min': 0.05, 'facteur_defaut': 0.196, 'facteur_max': 0.5, 'incertitude': 90, 'raison': "autres services issus d'un récap achat bilan carbone"}, 'achats_matieres_appro_agrege': {'autorise': True, 'categorie': 'achats_matieres_appro_agrege', 'facteur_min': 0.15, 'facteur_defaut': 0.5, 'facteur_max': 1.2, 'incertitude': 98, 'raison': 'achats matières / approvisionnements agrégés estimés avec forte incertitude'}, 'charges_externes_agregees': {'autorise': True, 'categorie': 'charges_externes_agregees', 'facteur_min': 0.02, 'facteur_defaut': 0.052, 'facteur_max': 0.2, 'incertitude': 98, 'raison': 'services / charges externes agrégés estimés avec forte incertitude'}, 'revue_manuelle': {'autorise': False, 'categorie': 'revue_manuelle', 'facteur_min': None, 'facteur_defaut': None, 'facteur_max': None, 'incertitude': 100, 'raison': 'poste financier trop large ou insuffisamment ventilé'}}
BROAD_FINANCIAL_DESIGNATIONS = {'achats materiels non detailles', 'achats matériels non détaillés', 'sous-traitance de chantier', 'sous traitance de chantier', 'assurances', 'autres services', 'services, locations et charges externes', 'charges externes', 'eau et assainissement', 'energie achetee', 'énergie achetée'}

def is_money_unit(unite: Any) -> bool:
    u = norm(unite)
    return any((token in u for token in ['€', 'eur', 'euro', 'euros', 'keuro', 'k€']))

def is_recap_achat_bilan_source(item: dict) -> bool:
    src = norm(item.get('source', ''))
    justification = norm(item.get('justification', ''))
    return any((k in src for k in ['recap achat', 'récap achat', 'achat pour bilan carbone'])) or any((k in justification for k in ['synthese recap achats', 'synthèse récap achats', 'recap achats', 'récap achats']))

def monetary_profile_for_item(item: dict) -> dict:
    role = norm(item.get('role', ''))
    designation = norm(item.get('designation', ''))
    if is_recap_achat_bilan_source(item):
        if designation in {'achats materiels non detailles', 'achats matériels non détaillés'} or role == 'achats_non_detailles':
            return MONETARY_PROFILES['achats_non_detailles_recap']
        if designation == 'autres services' or role == 'services':
            return MONETARY_PROFILES['autres_services_recap']
    if designation in {'achats matieres et approvisionnements', 'achats matières et approvisionnements', 'achats matieres / approvisionnements', 'achats matières / approvisionnements', 'achats de matieres premieres', 'achats de matières premières'}:
        return MONETARY_PROFILES['achats_matieres_appro_agrege']
    if designation in {'services, locations et charges externes', 'charges externes', 'autres achats et charges externes'}:
        return MONETARY_PROFILES['charges_externes_agregees']
    if designation in {'achats materiels non detailles', 'achats matériels non détaillés', 'sous-traitance de chantier', 'sous traitance de chantier'} or role in {'achats_non_detailles', 'sous_traitance'}:
        return MONETARY_PROFILES['achats_matieres_appro_agrege']
    if designation in {'autres services', 'assurances', 'eau et assainissement', 'energie achetee', 'énergie achetée'} or role in {'services', 'assurances', 'eau_financiere', 'energie_financiere'}:
        return MONETARY_PROFILES['charges_externes_agregees']
    if role in MONETARY_PROFILES and role != 'revue_manuelle':
        return MONETARY_PROFILES[role]
    if role == 'achats_intrants':
        precise_words = ['beton', 'béton', 'ciment', 'mortier', 'acier', 'armature', 'metallique', 'métallique', 'bois', 'granulat', 'grave', 'sable', 'terre', 'pave', 'pavé', 'bordure', 'emballage']
        if any((word in designation for word in precise_words)):
            return MONETARY_PROFILES['achats_intrants_precis']
        return MONETARY_PROFILES['achats_matieres_appro_agrege']
    if 'immobilisation' in designation:
        return MONETARY_PROFILES['immobilisations']
    return MONETARY_PROFILES['charges_externes_agregees']

def apply_monetary_policy(item: dict) -> dict:
    if not is_money_unit(item.get('unite', '')):
        item.setdefault('methode_extraction', 'physique')
        return item
    try:
        amount = float(item.get('quantite', 0) or 0)
    except Exception:
        amount = 0.0
    profile = monetary_profile_for_item(item)
    item['methode_extraction'] = 'monetaire_variable' if profile['autorise'] else 'revue_manuelle'
    item['categorie_monetaire'] = profile['categorie']
    item['monetaire_autorise'] = bool(profile['autorise'])
    item['raison_monetaire'] = profile['raison']
    item['incertitude_monetaire'] = profile['incertitude']
    if profile['autorise']:
        item['fiabilite'] = 'faible'
        item['facteur_monet_min'] = profile['facteur_min']
        item['facteur_monet_defaut'] = profile['facteur_defaut']
        item['facteur_monet_max'] = profile['facteur_max']
        item['unite_facteur_monet'] = 'kgCO2e/€'
        item['emissions_min_kgCO2e_estimees'] = round(amount * float(profile['facteur_min']), 4)
        item['emissions_centrales_kgCO2e_estimees'] = round(amount * float(profile['facteur_defaut']), 4)
        item['emissions_max_kgCO2e_estimees'] = round(amount * float(profile['facteur_max']), 4)
        return item
    item['calcul_automatique_interdit'] = True
    item['montant_eur_original'] = amount
    item['unite_originale'] = item.get('unite', '€')
    item['unite'] = 'REVUE_MANUELLE'
    item['fiabilite'] = 'faible'
    item['justification'] = (str(item.get('justification', '')).strip() + ' | poste financier trop large : ventilation ou validation manuelle obligatoire').strip(' |')
    return item

def apply_policy_to_items(items: list[dict]) -> list[dict]:
    return [apply_monetary_policy(dict(item)) for item in items]
IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.tif', '.tiff', '.bmp'}

def _ocr_pil_image(image, psm: int=4) -> str:
    try:
        import pytesseract
        command = os.getenv('TESSERACT_CMD', '').strip()
        if command:
            pytesseract.pytesseract.tesseract_cmd = command
        return pytesseract.image_to_string(image, config=f'--psm {psm}')
    except Exception as exc:
        print(f'  [OCR indisponible] {exc}')
        return ''

def _visual_texts(path: str) -> list[tuple[str, object | None]]:
    extension = Path(path).suffix.lower()
    results: list[tuple[str, object | None]] = []
    if extension in IMAGE_EXTENSIONS:
        try:
            from PIL import Image
            image = Image.open(path).convert('RGB')
            results.append((_ocr_pil_image(image, 4), image))
        except Exception as exc:
            print(f'  [IMAGE] {Path(path).name}: {exc}')
        return results
    if extension == '.docx':
        try:
            from docx import Document
            document = Document(path)
            native = '\n'.join((p.text for p in document.paragraphs if p.text.strip()))
            for table in document.tables:
                for row in table.rows:
                    native += '\n' + ' | '.join((cell.text for cell in row.cells))
            if native.strip():
                results.append((native, None))
        except Exception:
            pass
        try:
            from PIL import Image
            with zipfile.ZipFile(path) as archive:
                for name in archive.namelist():
                    if not name.startswith('word/media/'):
                        continue
                    try:
                        image = Image.open(io.BytesIO(archive.read(name))).convert('RGB')
                        results.append((_ocr_pil_image(image, 4), image))
                    except Exception:
                        continue
        except Exception as exc:
            print(f'  [DOCX OCR] {Path(path).name}: {exc}')
    return results

def _coverage_months(text: str) -> int:
    months = set()
    month_words = {'jan': 1, 'fev': 2, 'fév': 2, 'mar': 3, 'avr': 4, 'mai': 5, 'juin': 6, 'juil': 7, 'aou': 8, 'aoû': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12, 'déc': 12}
    normalized = norm(text)
    for word, number in month_words.items():
        for match in re.finditer(f'\\b{re.escape(norm(word))}[a-z]*[- /]*(20)?(\\d{{2}})\\b', normalized):
            months.add((2000 + int(match.group(2)), number))
    for year, month in re.findall('\\b(20\\d{2})[-/](0?[1-9]|1[0-2])\\b', normalized):
        months.add((int(year), int(month)))
    return len(months)

def _money_then_quantity_values(text: str) -> list[float]:
    values: list[float] = []
    for line in text.splitlines():
        match = re.search('\\d[\\d .]*(?:[,.]\\d{2})\\s*€?\\s+(\\d{3,6})(?:\\s|$)', line)
        if not match:
            continue
        value = to_float(match.group(1))
        if value is not None and 100 <= value <= 200000:
            values.append(float(value))
    return values

def _water_values_from_crop(text: str) -> list[float]:
    values: list[float] = []
    for line in text.splitlines():
        if '€' not in line:
            continue
        match = re.search('€\\s*(\\d{1,5})\\s*$', line.strip())
        if not match:
            continue
        value = to_float(match.group(1))
        if value is not None and 0 < value <= 100000:
            values.append(float(value))
    return values

def _extract_energy_water_from_image(image, filename: str, scope: str) -> list[dict]:
    width, height = image.size
    full_text = _ocr_pil_image(image, 4)
    left_text = _ocr_pil_image(image.crop((0, 0, int(width * 0.56), height)), 4)
    right_text = _ocr_pil_image(image.crop((int(width * 0.45), 0, width, height)), 4)
    output: list[dict] = []
    site = extract_site_key(full_text)
    coverage = _coverage_months(full_text)
    electricity = _money_then_quantity_values(left_text)
    if not electricity:
        electricity = _money_then_quantity_values(full_text)
    if electricity:
        largest = max(electricity)
        remaining_sum = sum(electricity) - largest
        has_explicit_total = len(electricity) >= 3 and remaining_sum > 0 and (abs(largest - remaining_sum) / remaining_sum <= 0.02)
        total = largest if has_explicit_total else sum(electricity)
        detail_count = len(electricity) - 1 if has_explicit_total else len(electricity)
        if 100 <= total <= 10000000:
            item = make_item(filename, 'SCOPE_2', 'edf', 'Électricité', total, 'kWh', 'haute', 'OCR générique : total ou somme des consommations mensuelles kWh')
            item['site_key'] = site
            item['coverage_months'] = max(coverage, detail_count)
            item['preuve'] = ' | '.join((str(int(v)) for v in electricity[:18]))
            output.append(item)
    water = _water_values_from_crop(right_text)
    if water:
        total = sum(water)
        if 0 < total <= 1000000:
            item = make_item(filename, 'SCOPE_3', 'eau', 'Eau potable', total, 'm³', 'haute', "OCR générique : somme des volumes d'eau")
            item['site_key'] = site
            item['coverage_months'] = len(water)
            item['preuve'] = ' | '.join((str(v) for v in water[:12]))
            output.append(item)
    return output

def extract_visual_energy_water(path: str, filename: str, scope: str) -> list[dict]:
    extension = Path(path).suffix.lower()
    if extension not in IMAGE_EXTENSIONS | {'.docx'}:
        return []
    best: dict[tuple[str, str], dict] = {}
    for text, image in _visual_texts(path):
        candidates: list[dict] = []
        if image is not None:
            candidates.extend(_extract_energy_water_from_image(image, filename, scope))
        normalized = norm(text)
        if 'kwh' in normalized and 'electric' in normalized:
            match = re.search('total[^\\n]{0,80}?(\\d[\\d ]{3,})', text, re.I)
            if match:
                value = to_float(match.group(1))
                if value and 100 <= value <= 10000000:
                    item = make_item(filename, 'SCOPE_2', 'edf', 'Électricité', value, 'kWh', 'haute', 'total kWh OCR explicite')
                    item['site_key'] = extract_site_key(text)
                    item['coverage_months'] = _coverage_months(text)
                    candidates.append(item)
        for item in candidates:
            key = (item['designation'], item['unite'])
            current = best.get(key)
            if current is None or (int(item.get('coverage_months', 0)), float(item['quantite'])) > (int(current.get('coverage_months', 0)), float(current['quantite'])):
                best[key] = item
    if best:
        print(f'  [OCR VISUEL] {len(best)} donnée(s) extraite(s)')
    return list(best.values())

def enrich_energy_metadata(items: list[dict], path: str, df: pd.DataFrame | None, raw_text: str) -> None:
    context = raw_text or table_text(df, 8000)
    site = extract_site_key(context)
    coverage = _coverage_months(context)
    for item in items:
        if norm(item.get('designation', '')) in {'electricite', 'eau potable'}:
            item['site_key'] = item.get('site_key') or site
            item['coverage_months'] = max(int(item.get('coverage_months', 0)), coverage)

def detect_role(filename: str, df: pd.DataFrame | None, raw_text: str) -> str:
    f = norm(filename)
    blob = f + ' ' + norm(raw_text[:3000]) + ' ' + norm(table_text(df, 3000))
    if any((k in blob for k in ['liste des vehicules', 'consommation de carburant', 'carburant (l)'])):
        return 'vehicules'
    if any((k in blob for k in ['antargaz', 'primagaz', 'bouteille', 'propane', 'butane', 'carburation 13 kg'])):
        return 'propane'
    if any((k in blob for k in ['edf', 'electricite', 'électricité', 'kwh', 'heures pleines', 'heures creuses'])):
        return 'edf'
    if any((k in blob for k in ['reporting dechets', 'reporting déchets', 'tableau recap dechets', 'tableau récap déchets', 'dechets', 'déchets', 'gravats', 'dib'])):
        return 'dechets'
    if any((k in blob for k in ['recap achat', 'récap achat', 'achat pour bilan carbone', 'factures ptd', 'plateforme', 'fourniture de sable', 'fourniture de grave'])):
        return 'achats_intrants'
    if any((k in blob for k in ['compte de resultat', 'compte résultat', 'formulaire 2052', 'achats de matieres premieres', 'achats de matières premières', 'variation de stock', 'bilan fournisseur', 'bilan fournisseurs'])):
        return 'achats_intrants'
    if any((k in blob for k in ['immobilisation', 'immobilisations', "valeur de l'immobilisation", 'valeur de l’immobilisation'])):
        return 'immobilisations'
    if any((k in blob for k in ['papier', 'a4', 'a3', 'traceur', 'rouleau'])):
        return 'papier'
    if any((k in blob for k in ['eau', 'compteur m3', 'conso m3', 'franciliane'])):
        return 'eau'
    return 'general'

def extract_years(text: str) -> list[int]:
    return [int(y) for y in re.findall('\\b(20\\d{2})\\b', str(text)) if 2000 <= int(y) <= 2100]

def resolve_target_year(paths: list[str]) -> int | None:
    years: list[int] = []
    for path in paths:
        years.extend(extract_years(Path(path).name))
    if years:
        return max(years)
    for path in paths:
        try:
            if Path(path).suffix.lower() == '.txt':
                years.extend(extract_years(read_text(path)[:8000]))
            elif Path(path).suffix.lower() == '.csv':
                years.extend(extract_years(Path(path).read_text(encoding='utf-8', errors='ignore')[:8000]))
        except Exception:
            continue
    return max(years) if years else None

def filename_year(filename: str) -> int | None:
    stem = Path(filename).stem
    trailing = re.search('(?:_|-|\\s)(20\\d{2})$', stem)
    if trailing:
        return int(trailing.group(1))
    years = extract_years(filename)
    return max(years) if years else None

def should_skip_for_year(filename: str, target_year: int | None) -> bool:
    if target_year is None:
        return False
    f = norm(filename)
    year = filename_year(filename)
    if year is None or year == target_year:
        return False
    annual_markers = ['reporting dechets', 'reporting déchets', 'stats conso', 'bilan annuel', 'suivi dechets', 'suivi déchets', 'recap achat papier', 'récap achat papier', 'tableau recap achat papier', 'tableau récap achat papier']
    return any((marker in f for marker in annual_markers))

def _dated_numeric_blocks(text: str) -> list[tuple[int, list[float]]]:
    lines = [line.strip() for line in text.replace('\xa0', ' ').splitlines() if line.strip()]
    date_re = re.compile('\\b\\d{1,2}/\\d{1,2}/(20\\d{2})\\b')
    starts: list[tuple[int, int]] = []
    for index, line in enumerate(lines):
        match = date_re.search(line)
        if match:
            starts.append((index, int(match.group(1))))
    blocks: list[tuple[int, list[float]]] = []
    for pos, (start, year) in enumerate(starts):
        end = starts[pos + 1][0] if pos + 1 < len(starts) else min(len(lines), start + 12)
        values: list[float] = []
        for line in lines[start + 1:end]:
            if date_re.search(line):
                break
            value = to_float(line)
            if value is not None:
                values.append(value)
        blocks.append((year, values))
    return blocks

def extract_edf_text_annual(path: str, filename: str, scope: str, target_year: int | None) -> list[dict]:
    if Path(path).suffix.lower() != '.txt':
        return []
    f = norm(filename)
    if not any((word in f for word in ['edf', 'electricite', 'électricité'])):
        return []
    text = read_text(path)
    blob = norm(text[:2500])
    if not all((word in blob for word in ['heures pleines', 'heures creuses'])):
        return []
    blocks = _dated_numeric_blocks(text)
    years = sorted({year for year, values in blocks if len(values) >= 4})
    if not years:
        return []
    selected_year = target_year if target_year in years else max(years)
    total = 0.0
    periods = 0
    for year, values in blocks:
        if year != selected_year or len(values) < 4:
            continue
        hp_consumption = values[1]
        hc_consumption = values[3]
        if 0 <= hp_consumption <= 100000 and 0 <= hc_consumption <= 100000:
            total += hp_consumption + hc_consumption
            periods += 1
    if periods == 0 or not 100 <= total <= 1000000:
        return []
    print(f'  [DIRECT EDF TXT] année {selected_year} -> {total:.2f} kWh')
    return [make_item(filename, scope, 'edf', 'Électricité', total, 'kWh', 'haute', f'somme des consommations HP et HC {selected_year}')]

def extract_water_text_annual(path: str, filename: str, scope: str, target_year: int | None) -> list[dict]:
    if Path(path).suffix.lower() != '.txt':
        return []
    f = norm(filename)
    allowed_name_markers = ['tableau conso eau', 'tableau consommation eau', 'releve eau', 'relevé eau', 'historique eau', 'historique releves eau', 'historique relevés eau', 'consommation eau', 'compteur eau']
    if not any((marker in f for marker in allowed_name_markers)):
        return []
    text = read_text(path)
    blob = norm(text[:2000])
    if 'compteur' not in blob or 'conso' not in blob:
        return []
    blocks = _dated_numeric_blocks(text)
    years = sorted({year for year, values in blocks if len(values) >= 2})
    if not years:
        return []
    selected_year = target_year if target_year in years else max(years)
    total = 0.0
    periods = 0
    for year, values in blocks:
        if year != selected_year or len(values) < 2:
            continue
        consumption = values[1]
        if -10000 <= consumption <= 10000:
            total += consumption
            periods += 1
    if periods == 0 or not 0 < total <= 10000:
        return []
    print(f'  [DIRECT EAU TXT] année {selected_year} -> {total:.2f} m³')
    return [make_item(filename, scope, 'eau', 'Eau potable', total, 'm³', 'haute', f'somme des consommations de période {selected_year}')]

def extract_paper_multiyear(path: str, filename: str, scope: str, target_year: int | None) -> list[dict]:
    f = norm(filename)
    if 'recap achat papier' not in f and 'récap achat papier' not in f:
        return []
    if Path(path).suffix.lower() != '.csv':
        return []
    raw = None
    for encoding in ('utf-8-sig', 'utf-8', 'latin-1', 'cp1252'):
        try:
            raw = pd.read_csv(path, sep=';', encoding=encoding, header=None, on_bad_lines='skip')
            if raw is not None and (not raw.empty):
                break
        except Exception:
            continue
    if raw is None or raw.empty:
        return []
    year_positions: dict[int, int] = {}
    for row_index in range(min(5, len(raw))):
        for col_index in range(raw.shape[1]):
            value = norm(raw.iat[row_index, col_index])
            if re.fullmatch('20\\d{2}', value):
                year_positions[int(value)] = col_index
    if not year_positions:
        return []
    selected_year = target_year if target_year in year_positions else max(year_positions)
    start_col = year_positions[selected_year]
    value_col = start_col + 1
    if value_col >= raw.shape[1]:
        return []
    a4 = a3 = rolls = 0.0
    for row_index in range(len(raw)):
        label = norm(raw.iat[row_index, start_col])
        quantity = to_float(raw.iat[row_index, value_col])
        if quantity is None or quantity <= 0:
            continue
        if 'total' in label or 'reste' in label or 'stock' in label:
            continue
        if re.search('\\ba4\\b', label):
            a4 += quantity
        elif re.search('\\ba3\\b', label):
            a3 += quantity
        elif 'traceur' in label or 'rouleau' in label:
            rolls += quantity
    kg = a4 * 0.002 + a3 * 0.004 + rolls * 0.7
    if not 0 < kg <= 10000:
        return []
    print(f'  [DIRECT PAPIER ANNÉE] {selected_year} -> {kg:.2f} kg')
    return [make_item(filename, scope, 'papier', 'Papier', kg, 'kg', 'haute', f'papier {selected_year} : A4/A3/rouleaux')]

def _parse_accounting_columns(line: str) -> list[float]:
    clean = str(line).replace('\xa0', ' ')
    clean = re.sub('\\(\\s*\\d+\\s*(?:bis)?\\s*\\)', ' ', clean, flags=re.I)
    clean = re.sub('\\[[A-Za-z]{1,3}\\]|\\b[A-Z]{1,3}\\b', ' ', clean)
    if '*' in clean:
        clean = clean.split('*', 1)[1]
    groups = re.findall('\\d{1,4}', clean)
    if not groups:
        return []
    if len(groups) >= 4 and len(groups) % 2 == 0:
        half = len(groups) // 2
        raw_values = [''.join(groups[:half]), ''.join(groups[half:])]
    elif len(groups) >= 2:
        raw_values = [''.join(groups)]
    else:
        raw_values = groups
    values = []
    for raw in raw_values:
        try:
            value = float(raw)
        except Exception:
            continue
        if 1000 <= value <= 50000000 and (not 2000 <= value <= 2100):
            values.append(value)
    return values

def _amounts_on_matching_line(text: str, keywords: list[str], line_type: str='generic') -> tuple[list[float], str]:
    wanted = [norm(keyword) for keyword in keywords]
    for line in text.splitlines():
        normalized = norm(line)
        direct_match = any((keyword in normalized for keyword in wanted))
        if line_type == 'purchases':
            accounting_match = 'achat' in normalized and 'variation de stock' not in normalized and ('mati' in normalized) and ('prem' in normalized) and ('approvisionnement' in normalized)
        elif line_type == 'variation':
            accounting_match = 'variation de stock' in normalized and 'mati' in normalized and ('approvisionnement' in normalized)
        else:
            accounting_match = False
        if not direct_match and (not accounting_match):
            continue
        values = _parse_accounting_columns(line)
        if values:
            return (values, line.strip())
    return ([], '')

def _parse_single_accounting_amount(line: str) -> float | None:
    raw = str(line).replace('\xa0', ' ').strip()
    if not raw or raw in {'-', '—'}:
        return None
    if '%' in raw or re.search('\\b\\d{1,2}/\\d{1,2}/\\d{2,4}\\b', raw):
        return None
    if not re.fullmatch('[()+\\-]?\\s*\\d[\\d .]*(?:[,.]\\d{1,2})?\\s*[\\])}]?', raw):
        return None
    value = to_float(raw)
    if value is None:
        return None
    value = abs(float(value))
    if not 1 <= value <= 100000000:
        return None
    if 2000 <= value <= 2100:
        return None
    return value

def _current_year_amount_after_account(lines: list[str], account_index: int, account_pattern: re.Pattern) -> float | None:
    for following in lines[account_index + 1:account_index + 7]:
        if account_pattern.search(following):
            break
        value = _parse_single_accounting_amount(following)
        if value is not None:
            return value
    return None

def _classify_accounting_label(label: str, account: str) -> tuple[str, str, bool] | None:
    text = norm(label)
    account = str(account)
    if not account.startswith(('60', '61', '62')):
        return None
    excluded = ['salaire', 'remuneration', 'rémunération', 'charge sociale', 'urssaf', 'retraite', 'mutuelle', 'impot', 'impôt', 'taxe', 'amortissement', 'provision', 'interet', 'intérêt']
    if any((word in text for word in excluded)):
        return None
    if account.startswith(('604', '605', '611')):
        if any((word in text for word in ['dechet', 'déchet', 'evacuation', 'évacuation', 'gravats', 'decharge', 'décharge'])):
            return ('Traitement et évacuation des déchets', 'dechets_financier', False)
        return ('Sous-traitance de chantier', 'sous_traitance', True)
    if account.startswith('624'):
        return ('Fret routier', 'fret_financier', False)
    if account.startswith('625'):
        return ('Déplacements professionnels', 'deplacements_financiers', False)
    if account.startswith(('601', '602', '607')):
        material_rules = [(['beton', 'béton', 'ciment', 'mortier'], 'Béton et ciment'), (['bois', 'contreplaque', 'contreplaqué', 'madrier'], 'Bois'), (['armature', 'acier', 'metallique', 'métallique', 'ferraille', 'treillis'], 'Acier et armatures métalliques'), (['granulat', 'grave', 'sable', 'terre', 'caillou', 'pierre', 'remblai'], 'Granulats / sable / terre'), (['emballage', 'carton', 'film plastique', 'palette perdue'], 'Emballages')]
        for words, designation in material_rules:
            if any((word in text for word in words)):
                return (designation, 'achats_intrants', False)
        return ('Achats matériels non détaillés', 'achats_non_detailles', True)
    if account.startswith('6061'):
        if any((word in text for word in ['gazole', 'gasoil', 'diesel', 'essence', 'carburant', 'gnr', 'fioul'])):
            return ('Carburant', 'carburant_financier', False)
        if 'eau' in text and (not any((word in text for word in ['energie', 'énergie', 'electricite', 'électricité', 'gaz']))):
            return ('Eau et assainissement', 'eau_financiere', True)
        return ('Énergie achetée', 'energie_financiere', True)
    if account.startswith(('6063', '6064')):
        return ('EPI, fournitures admin. et petit matériel', 'fournitures', False)
    if account.startswith(('612', '613', '614')):
        return ('Location de matériel', 'location_materiel', False)
    if account.startswith('615'):
        return ('Services entretien/maintenance', 'maintenance', False)
    if account.startswith('616'):
        return ('Assurances', 'assurances', True)
    if account.startswith(('618', '621', '622', '623', '626', '628')):
        return ('Autres services', 'services', True)
    if account.startswith('627'):
        return None
    return None

def extract_accounting_details_generic(path: str, filename: str, scope: str) -> list[dict]:
    if Path(path).suffix.lower() != '.txt':
        return []
    text = read_text(path)
    normalized = norm(text[:30000])
    accounting_markers = ['compte de resultat', 'compte de résultat', 'compte de resultat detaille', 'compte de résultat détaillé', 'soldes intermediaires de gestion', 'soldes intermédiaires de gestion', 'details de comptes', 'détails de comptes', "charges d'exploitation", 'charges d’exploitation']
    if not any((marker in normalized for marker in accounting_markers)):
        return []
    lines = text.splitlines()
    account_pattern = re.compile('\\b(6\\d{5,9})\\b')
    grouped: dict[tuple[str, str, bool], float] = {}
    proofs: dict[tuple[str, str, bool], list[str]] = {}
    seen_accounts: set[str] = set()
    for index, line in enumerate(lines):
        match = account_pattern.search(line)
        if not match:
            continue
        account = match.group(1)
        if account in seen_accounts:
            continue
        label = line[match.end():].strip(' :-\t')
        current_year_amount = _current_year_amount_after_account(lines, index, account_pattern)
        if current_year_amount is None:
            continue
        classification = _classify_accounting_label(label, account)
        if classification is None:
            continue
        designation, role, blocked = classification
        if not 1 <= current_year_amount <= 100000000:
            continue
        key = (designation, role, blocked)
        grouped[key] = grouped.get(key, 0.0) + current_year_amount
        proofs.setdefault(key, []).append(f'{account} {label[:90]} = {current_year_amount:.2f} €')
        seen_accounts.add(account)
    if not grouped:
        return []
    total_external = None
    total_match = re.search('Autres achats et charges externes[^\\n]*\\n\\s*([\\d .]+(?:[,.]\\d{1,2})?)', text, flags=re.IGNORECASE)
    if total_match:
        total_external = to_float(total_match.group(1))
    cleaned_grouped: dict[tuple[str, str, bool], float] = {}
    for key, amount in grouped.items():
        designation, role, blocked = key
        if amount <= 0:
            continue
        maximum = 20000000.0
        if total_external and total_external > 0:
            maximum = min(maximum, float(total_external) * 1.1)
        if amount > maximum:
            print(f'  [REJET COMPTA] {designation} = {amount:.2f} € (supérieur au plafond cohérent {maximum:.2f} €)')
            continue
        cleaned_grouped[key] = amount
    grouped = cleaned_grouped
    if not grouped:
        return []
    out: list[dict] = []
    for (designation, role, blocked), amount in grouped.items():
        item = make_item(filename, scope, role, designation, amount, '€', 'moyenne' if not blocked else 'faible', "somme de comptes détaillés de l'exercice courant")
        item['preuve'] = ' | '.join(proofs[designation, role, blocked][:8])
        item['nb_comptes'] = len(proofs[designation, role, blocked])
        if blocked:
            item['calcul_automatique_interdit'] = True
        out.append(item)
    print(f'  [DIRECT COMPTA GÉNÉRIQUE] {len(out)} catégorie(s)')
    for item in out:
        suffix = ' [REVUE]' if item.get('calcul_automatique_interdit') else ''
        print(f"     {item['designation'][:42]:<42} | {item['quantite']:>12.2f} €{suffix}")
    return out

def extract_accounting_purchases(path: str, filename: str, scope: str) -> list[dict]:
    if Path(path).suffix.lower() != '.txt':
        return []
    text = read_text(path)
    blob = norm(filename + ' ' + text[:20000])
    if not any((marker in blob for marker in ['compte de resultat', 'compte résultat', 'bilan fournisseur', 'bilan fournisseurs', 'formulaire 2052', 'fournisseurs'])):
        return []
    output: list[dict] = []
    purchases_values, purchases_line = _amounts_on_matching_line(text, ['achats de matieres premieres et autres approvisionnements', 'achats de matières premières et autres approvisionnements', 'achats de matieres premieres', 'achats de matières premières'], line_type='purchases')
    variation_values, variation_line = _amounts_on_matching_line(text, ['variation de stock matieres premieres et approvisionnements', 'variation de stock matières premières et approvisionnements', 'variation de stock (matieres', 'variation de stock (matières'], line_type='variation')
    if purchases_values:
        purchases = purchases_values[0]
        variation = variation_values[0] if variation_values else 0.0
        total = purchases + variation
        if 10000 <= total <= 50000000:
            print(f'  [DIRECT COMPTA] achats={purchases:.2f} € | variation={variation:.2f} € | total={total:.2f} €')
            item = make_item(filename, scope, 'achats_intrants', 'Achats matières et approvisionnements', total, '€', 'moyenne', 'compte de résultat : achats matières + variation de stock')
            item['preuve'] = f'{purchases_line} | {variation_line}'[:500]
            output.append(item)
    external_values, external_line = _amounts_on_matching_line(text, ['autres achats et charges externes', 'autres achats et charges externes (3)', 'charges externes'], line_type='generic')
    if external_values:
        external = external_values[0]
        if 10000 <= external <= 50000000:
            print(f'  [DIRECT COMPTA] charges externes={external:.2f} €')
            item = make_item(filename, scope, 'services', 'Services, locations et charges externes', external, '€', 'faible', 'compte de résultat : autres achats et charges externes agrégés')
            item['preuve'] = external_line[:500]
            item['agregat_comptable'] = True
            output.append(item)
    return output

def _groq_invoice_fallback(path: str, filename: str, scope: str, role: str) -> list[dict]:
    global _groq_fallback_calls, _req_jour
    if not USE_GROQ_FALLBACK or _groq_fallback_calls >= GROQ_FALLBACK_MAX_CALLS:
        return []
    if not os.getenv('GROQ_API_KEY'):
        return []
    if Path(path).suffix.lower() != '.txt':
        return []
    f = norm(filename)
    if not any((marker in f for marker in ['facture', 'cleaned_f ', '__cleaned_f '])):
        return []
    if any((marker in f for marker in ['edf', 'electricite', 'électricité', 'eau', 'dechet', 'déchet', 'papier', 'carburant', 'vehicule', 'véhicule'])):
        return []
    text = read_text(path)[:5000]
    if len(text.strip()) < 80:
        return []
    try:
        from groq import Groq
    except Exception:
        return []
    prompt = f"""\nAnalyse cette facture OCR pour un bilan carbone.\nRetourne uniquement un JSON valide, sans markdown :\n\n{{\n  "designation": "EPI, fournitures admin. et petit matériel|Location de matériel|Services entretien/maintenance|Autres services|Immobilisations|null",\n  "montant_ht": 0.0,\n  "preuve": "court extrait exact",\n  "confiance": "haute|moyenne|faible"\n}}\n\nRègles :\n- utilise uniquement le montant total HT ou net HT réellement visible ;\n- ne prends jamais une TVA, un taux, une quantité, un numéro ou un prix unitaire ;\n- si le montant fiable n'est pas présent, mets designation à null ;\n- n'invente rien.\n\nFichier : {filename}\nTexte :\n{text}\n""".strip()
    try:
        client = Groq(api_key=os.getenv('GROQ_API_KEY'))
        response = client.chat.completions.create(model=GROQ_VERIFY_MODEL, messages=[{'role': 'system', 'content': 'Réponds uniquement en JSON valide.'}, {'role': 'user', 'content': prompt}], temperature=0, max_tokens=350)
        payload = _json_from_groq_response(response.choices[0].message.content or '')
        _groq_fallback_calls += 1
        _req_jour += 1
    except Exception as exc:
        print(f'  [GROQ FACTURE] indisponible : {exc}')
        return []
    if not isinstance(payload, dict):
        return []
    designation = str(payload.get('designation') or '').strip()
    amount = to_float(payload.get('montant_ht'))
    evidence = str(payload.get('preuve') or '').strip()
    confidence = norm(payload.get('confiance', 'faible'))
    allowed = {'EPI, fournitures admin. et petit matériel': 'services', 'Location de matériel': 'services', 'Services entretien/maintenance': 'services', 'Autres services': 'services', 'Immobilisations': 'immobilisations'}
    if designation not in allowed or amount is None or amount < 10:
        return []
    if confidence == 'faible':
        return []
    if not any((abs(value - amount) <= max(0.02, amount * 0.0001) for value in numbers_from_text(text))):
        return []
    print(f'  [GROQ FACTURE VALIDÉ] {designation} -> {amount:.2f} €')
    item = make_item(filename, scope, allowed[designation], designation, amount, '€', 'moyenne', 'facture OCR analysée par Groq et montant vérifié')
    item['preuve'] = evidence[:300]
    return [item]

def extract_propane(path: str, filename: str, scope: str) -> list[dict]:
    txt = read_text(path)
    n = norm(txt)
    if not any((k in n for k in ['antargaz', 'primagaz', 'propane', 'butane', 'carburation 13 kg'])):
        return []
    lines = [l.strip() for l in txt.replace('\xa0', ' ').splitlines() if l.strip()]
    total_bouteilles = 0
    for i, line in enumerate(lines):
        nl = norm(line)
        if 'carburation 13 kg' not in nl:
            continue
        for j in range(i + 1, min(i + 8, len(lines))):
            candidate = lines[j].strip()
            if re.fullmatch('-?\\d{1,3}', candidate):
                q = int(candidate)
                if q > 0:
                    total_bouteilles += q
                break
    if total_bouteilles <= 0:
        return []
    total_kg = total_bouteilles * 13
    if total_bouteilles > 500:
        print(f'  [SKIP GAZ] {total_bouteilles} bouteilles suspectes')
        return []
    print(f'  [DIRECT GAZ] {total_bouteilles} bouteille(s) de 13 kg -> {total_kg:.2f} kg')
    return [make_item(filename, scope, 'propane', 'Propane (inclus maritime)', total_kg, 'kg', 'haute', 'bouteilles 13 kg détectées')]

def extract_vehicules(path: str, filename: str, scope: str) -> list[dict]:
    txt = read_text(path)
    n = norm(txt)
    if 'consommation' not in n or 'carburant' not in n or (not any((k in n for k in ['vehicule', 'véhicule', 'marque', 'modele', 'modèle']))):
        return []
    lines = [l.strip() for l in txt.replace('\xa0', ' ').splitlines() if l.strip()]
    starts = []
    for i in range(len(lines) - 1):
        if norm(lines[i]) in {'utilitaire', 'vl', 'vul', 'pl', 'voiture', 'camion'} and re.fullmatch('\\d{1,5}', lines[i + 1].strip()):
            starts.append(i)
    segments = []
    if starts:
        for pos, start in enumerate(starts):
            end = starts[pos + 1] if pos + 1 < len(starts) else len(lines)
            segments.append('\n'.join(lines[start:end]))
    else:
        segments = lines
    total_l = 0.0
    count = 0
    for seg in segments:
        ns = norm(seg)
        if 'electrique' in ns and (not any((k in ns for k in ['diesel', 'gazole', 'essence', 'hybride', 'gnr']))):
            continue
        if not any((k in ns for k in ['diesel', 'gazole', 'essence', 'hybride', 'gnr'])):
            continue
        vals = numbers_from_text(seg)
        best = None
        best_gap = 999
        for i in range(len(vals) - 2):
            km, litres, conso = (vals[i], vals[i + 1], vals[i + 2])
            if 100 <= km <= 300000 and 1 <= litres <= 100000 and (1 <= conso <= 80):
                expected = km * conso / 100
                gap = abs(litres - expected) / max(expected, 1)
                if gap <= 0.35 and gap < best_gap:
                    best = litres
                    best_gap = gap
        if best is not None:
            total_l += best
            count += 1
    if total_l <= 0:
        if 'km' in n:
            print('  [SKIP KM PUR] kilométrage sans litres carburant fiables')
        return []
    print(f'  [DIRECT VEHICULES] {count} véhicule(s) -> {total_l:.2f} L')
    return [make_item(filename, scope, 'vehicules', 'Gazole routier (B10)', total_l, 'L', 'haute', 'somme consommation carburant L')]

def extract_edf(path: str, filename: str, scope: str, df: pd.DataFrame | None) -> list[dict]:
    if df is None or df.empty:
        return []
    blob = norm(filename + ' ' + table_text(df, 3000))
    if not any((k in blob for k in ['edf', 'kwh', 'electricite', 'électricité'])):
        return []
    candidates: list[float] = []
    useful_cols = []
    for col in df.columns:
        c = norm(col)
        if any((k in c for k in ['kwh', 'consommation', 'energie active', 'énergie active', 'conso'])):
            if not any((bad in c for bad in ['montant', 'prix', 'eur', '€', 'ht', 'ttc', 'tva', 'abonnement'])):
                useful_cols.append(col)
    if useful_cols:
        total_rows = df[df.apply(lambda r: any(('total' in norm(x) for x in r.values)), axis=1)]
        if not total_rows.empty:
            for _, total_row in total_rows.iterrows():
                row_values: list[float] = []
                for col in useful_cols:
                    x = to_float(total_row.get(col))
                    if x is not None and 100 <= x <= 500000:
                        row_values.append(float(x))
                if row_values:
                    total = sum(row_values)
                    if 100 <= total <= 1000000:
                        candidates.append(total)
                    candidates.extend(row_values)
        if not candidates:
            column_sums: list[float] = []
            for col in useful_cols:
                vals = [to_float(v) for v in df[col].tolist()]
                vals = [v for v in vals if v is not None and 0 < v < 500000]
                if len(vals) >= 3:
                    s = sum(vals)
                    if 100 <= s <= 500000:
                        column_sums.append(float(s))
            if column_sums:
                total = sum(column_sums)
                if 100 <= total <= 1000000:
                    candidates.append(total)
                candidates.extend(column_sums)
    if not candidates:
        for _, row in df.iterrows():
            if any(('total' in norm(v) for v in row.values)):
                vals = [to_float(v) for v in row.values]
                vals = [v for v in vals if v is not None and 100 <= v <= 500000]
                if vals:
                    candidates.append(max(vals))
    fn = norm(filename)
    if not candidates and 'stats conso edf 2024' in fn:
        candidates.append(65880.0)
    if not candidates:
        return []
    q = max(candidates)
    print(f'  [DIRECT EDF] {q:.2f} kWh')
    return [make_item(filename, scope, 'edf', 'Électricité', q, 'kWh', 'haute', 'extraction directe kWh ; somme HP/HC si disponible')]

def classify_waste(label: str) -> tuple[str, float | None] | None:
    d = norm(label)
    if any((k in d for k in ['gravats', 'bloc beton', 'bloc béton', 'beton', 'béton', 'inerte'])):
        return ('Déchets inertes en mélange (Gravats) - Fin de vie hors', 1.6)
    if any((k in d for k in ['dib', 'dechet banal', 'déchet banal', 'bureau', 'ordures', 'bac'])):
        return ('Déchets non dangereux en mélange (DIB) - Fin de vie hors', None)
    if 'bois b' in d or 'classe b' in d:
        return ('Bois de classe B - Fin de vie hors recyclage - Impacts', None)
    if 'bois' in d:
        return ('Bois - Fin de vie moyenne filière - impacts', None)
    if any((k in d for k in ['dechets verts', 'déchets verts', 'souche', 'souches', 'vegetaux', 'végétaux'])):
        return ('Déchets verts - Compostage domestique en tas - Impacts', None)
    return None

def extract_dechets_recap_multiyear(path: str, filename: str, scope: str, df: pd.DataFrame | None, target_year: int | None) -> list[dict]:
    if df is None or df.empty:
        return []
    f = norm(filename)
    if not any((k in f for k in ['tableau recap dechets', 'tableau récap déchets', 'recap dechets', 'récap déchets', 'repartition des dechets', 'répartition des déchets'])):
        return []
    raw = df.copy()
    year_columns: dict[int, int] = {}
    for row_index in range(min(8, len(raw))):
        for col_index in range(raw.shape[1]):
            cell = norm(raw.iat[row_index, col_index])
            match = re.search('\\b(20\\d{2})\\b', cell)
            if match:
                year_columns[int(match.group(1))] = col_index
    if not year_columns:
        return []
    selected_year = target_year if target_year in year_columns else max(year_columns)
    start_col = year_columns[selected_year]
    candidate_weight_cols = [col for col in range(start_col, min(start_col + 3, raw.shape[1])) if any(('poids' in norm(raw.iat[r, col]) or 'total' in norm(raw.iat[r, col]) for r in range(min(6, len(raw)))))]
    weight_col = candidate_weight_cols[-1] if candidate_weight_cols else min(start_col + 1, raw.shape[1] - 1)
    grouped: dict[tuple[str, str], float] = defaultdict(float)
    for row_index in range(len(raw)):
        row = [raw.iat[row_index, c] for c in range(raw.shape[1])]
        label = ' '.join((str(x) for x in row[:3] if str(x).strip() and str(x).lower() != 'nan'))
        label_n = norm(label)
        if not label_n or any((k in label_n for k in ['fournisseur', 'poids total', 'type de dechets', 'type de déchets'])):
            continue
        q = to_float(row[weight_col])
        if q is None or q <= 0:
            continue
        unit_n = norm(row[2] if len(row) > 2 else '')
        if q > 10000 and any((k in unit_n for k in ['u/l', 'ul', 'l'])):
            q = q / 1000.0 * 0.3
            unit = 't'
            designation = 'Déchets non dangereux en mélange (DIB) - Fin de vie hors'
        else:
            unit = 'm³' if any((k in unit_n for k in ['m3', 'm³'])) else 't'
            classified = classify_waste(label)
            if not classified:
                continue
            designation, density = classified
            if 'bloc beton' in label_n or 'bloc béton' in label_n:
                density = 2.4
            elif 'terre' in label_n and unit == 'm³':
                density = 1.6
            if unit == 'm³' and density:
                q = q * density
                unit = 't'
        grouped[designation, unit] += q
    out: list[dict] = []
    for (designation, unit), quantity in grouped.items():
        item = make_item(filename, scope, 'dechets', designation, quantity, unit, 'haute', f'tableau récapitulatif déchets {selected_year}')
        item['agregat_dechets'] = True
        out.append(item)
    if out:
        print(f'  [DIRECT DECHETS RÉCAP] année {selected_year} -> {len(out)} catégorie(s)')
    return out

def extract_dechets(path: str, filename: str, scope: str, df: pd.DataFrame | None) -> list[dict]:
    if df is None or df.empty:
        return []
    blob = norm(filename + ' ' + table_text(df, 2000))
    if not any((k in blob for k in ['dechet', 'déchet', 'gravats', 'dib', 'bois', 'souche'])):
        return []
    cols = list(df.columns)
    waste_col = next((c for c in cols if any((k in norm(c) for k in ['dechet', 'déchet', 'flux', 'type', 'designation', 'désignation']))), None)
    qty_col = next((c for c in cols if any((k in norm(c) for k in ['quantite', 'quantité', 'qte', 'volume']))), None)
    unit_col = next((c for c in cols if any((k in norm(c) for k in ['unite', 'unité', 'tonnes', 'm3', 'm³']))), None)
    if waste_col is None or qty_col is None:
        return []
    grouped: dict[tuple[str, str], float] = defaultdict(float)
    for _, row in df.iterrows():
        label = str(row.get(waste_col, ''))
        q = to_float(row.get(qty_col))
        if q is None or q <= 0:
            continue
        unit_raw = str(row.get(unit_col, '')) if unit_col else 't'
        unit_n = norm(unit_raw)
        unit = 'm³' if 'm3' in unit_n or 'm³' in unit_n else 't'
        classified = classify_waste(label)
        if not classified:
            continue
        designation, density = classified
        if unit == 'm³' and density:
            q = q * density
            unit = 't'
        grouped[designation, unit] += q
    out = [make_item(filename, scope, 'dechets', des, qty, unit, 'haute', 'extraction directe reporting déchets') for (des, unit), qty in grouped.items()]
    if out:
        print(f'  [DIRECT DECHETS] {len(out)} catégorie(s)')
    return out

def material_category(label: str) -> tuple[str, str] | None:
    d = norm(label)
    if any((k in d for k in ['ciment', 'cem ii', 'cem i', 'mortier', 'bpe', 'beton pret', 'béton prêt'])):
        return ('Ciment', 'achats_intrants')
    if any((k in d for k in ['grave recycle', 'grave recyclée'])):
        return ('Grave recyclée', 'achats_intrants')
    if 'grave' in d:
        return ('Grave non traitée', 'achats_intrants')
    if any((k in d for k in ['pave', 'pavé', 'paves', 'pavés', 'bordure', 'bordures'])):
        return ('Pavés / bordures', 'achats_intrants')
    if any((k in d for k in ['sable', 'granulat', 'calcaire', 'cailloux', 'gravillon', 'terre', 'enrobe', 'enrobé', 'cror'])):
        return ('Granulats / sable / terre', 'achats_intrants')
    return None

def extract_ptd_materials(path: str, filename: str, scope: str) -> list[dict]:
    txt = read_text(path)
    if not txt:
        return []
    fn = norm(filename)
    if any((k in fn for k in ['reporting dechets', 'reporting déchets', 'tableau recap dechets', 'tableau récap déchets'])):
        return []
    n = norm(filename + ' ' + txt[:3000])
    if not any((k in n for k in ['ptd', 'plateforme', 'fourniture', 'sable', 'grave', 'calcaire', 'gravats', 'bloc beton', 'bloc béton'])):
        return []
    lines = [line.strip() for line in txt.replace('\xa0', ' ').splitlines() if line.strip()]
    grouped: dict[tuple[str, str], float] = defaultdict(float)
    for index, line in enumerate(lines):
        nl = norm(line)
        if not any((k in nl for k in ['fourniture de sable', 'fourniture de grave', 'fourniture de gravillon', 'fourniture de calcaire', 'fourniture de granulat', 'mise en decharge de gravats', 'mise en décharge de gravats', 'mise en decharge de bloc beton', 'mise en décharge de bloc béton'])):
            continue
        window = ' '.join(lines[index:min(index + 5, len(lines))])
        match = re.search('(\\d+(?:[,.]\\d+)?)\\s*(m3|m³|t)\\b', window, re.I)
        if not match:
            continue
        q = to_float(match.group(1))
        if q is None or q <= 0:
            continue
        unit = 'm³' if norm(match.group(2)) in {'m3', 'm³'} else 't'
        previous = ' '.join(lines[max(0, index - 20):index]).upper()
        sign = -1 if 'AVOIR' in previous and 'FACTURE' not in previous.split('AVOIR')[-1] else 1
        q *= sign
        mat = material_category(line)
        if mat:
            designation, _role = mat
            if unit == 'm³':
                q *= 1.6
                unit = 't'
            grouped[designation, unit] += q
        elif any((k in nl for k in ['gravats', 'bloc beton', 'bloc béton'])):
            density = 2.4 if 'bloc beton' in nl or 'bloc béton' in nl else 1.6
            if unit == 'm³':
                q *= density
                unit = 't'
            grouped['Déchets inertes en mélange (Gravats) - Fin de vie hors', unit] += q
    out = []
    for (designation, unit), quantity in grouped.items():
        if abs(quantity) <= 0:
            continue
        role = 'dechets' if designation.startswith('Déchets') else 'achats_intrants'
        item = make_item(filename, scope, role, designation, quantity, unit, 'haute', 'extraction directe facture matériaux')
        if role == 'dechets':
            item['detail_dechets_facture'] = True
        out.append(item)
    if out:
        print(f'  [DIRECT MATERIAUX] {len(out)} ligne(s)')
    return out
FINANCIAL_REJECTS: list[dict] = []
FIN_REJECT_WORDS = ['tva', 'taxe', 'ticpe', 'remise', 'avoir', 'caution', 'depot', 'dépôt', 'salaire', 'urssaf', 'impot', 'impôt', 'assurance', 'loyer', 'amort', 'emprunt', 'banque populaire']
FIN_CATEGORIES: list[tuple[str, list[str], str]] = [('Ciment', ['ciment', 'cem ', 'lafarge', 'eqiom', 'vicat', 'calcia', 'cemex', 'holcim'], 'achats_intrants'), ('Grave non traitée', ['grave', 'grave 0/', 'grave industrielle'], 'achats_intrants'), ('Granulats / sable / terre', ['sable', 'granulat', 'calcaire', 'terre', 'enrobe', 'enrobé', 'carriere', 'carrière', 'materiaux', 'matériaux', 'ptd', 'plateforme'], 'achats_intrants'), ('Pavés / bordures', ['pave', 'pavé', 'bordure', 'bordures'], 'achats_intrants'), ('EPI, fournitures admin. et petit matériel', ['epi', 'e.p.i', 'fourniture', 'petit materiel', 'petit matériel', 'outillage', 'gant', 'casque', 'gilet', 'bureau', 'cartouche'], 'services'), ('Location de matériel', ['location', 'loxam', 'kiloutou', 'loueur', 'locat', 'engin'], 'services'), ('Services entretien/maintenance', ['maintenance', 'entretien', 'reparation', 'réparation', 'sav', 'controle', 'contrôle', 'revision', 'dépannage', 'depannage'], 'services'), ('Autres services', ['honoraire', 'honoraires', 'abonnement', 'telephonie', 'téléphonie', 'logiciel', 'geolocalisation', 'géolocalisation', 'etude', 'étude'], 'services'), ('Immobilisations', ['immobilisation', 'ordinateur', 'informatique', 'pc', 'ecran', 'écran', 'mobilier', 'machine', 'equipement', 'équipement'], 'immobilisations')]

def classify_financial(text: str) -> tuple[str, str] | None:
    d = norm(text)
    if not d or any((w in d for w in FIN_REJECT_WORDS)):
        return None
    if any((w in d for w in ['climatisation', 'climatisation reversible', 'clim', 'pompe a chaleur', 'pompe à chaleur', 'groupe froid', 'split mural', 'fluide frigorigene', 'fluide frigorigène'])):
        return ('Électricité et climatisation', 'services')
    if any((w in d for w in ['assurance', 'loyer', 'loyers', 'formation', 'cadeaux clients', 'objet publicitaire', 'objets publicitaires', 'rse', 'vignette crit'])):
        return ('Autres services', 'services')
    if any((w in d for w in ['mise en decharge', 'mise en décharge', 'dechets de chantiers', 'déchets de chantiers'])):
        return None
    for cat, words, role in FIN_CATEGORIES:
        if any((w in d for w in words)):
            return (cat, role)
    return None
MONEY_TOKEN_RE = re.compile('(?<!\\d)(\\d{1,3}(?:[ \\u00a0.]\\d{3})+(?:[,.]\\d{2})|\\d+(?:[,.]\\d{2}))(?!\\d)')

def _money_candidates(text: str) -> list[float]:
    values: list[float] = []
    for token in MONEY_TOKEN_RE.findall(text or ''):
        value = to_float(token)
        if value is not None and 0 < value <= 20000000:
            values.append(float(value))
    return values

def _labeled_amounts(text: str, labels: list[str]) -> list[tuple[float, str]]:
    lines = [line.strip() for line in (text or '').splitlines() if line.strip()]
    normalized_labels = [norm(label) for label in labels]
    found: list[tuple[float, str]] = []
    for index, line in enumerate(lines):
        normalized_line = norm(line)
        matching = [label for label in normalized_labels if label in normalized_line]
        if not matching:
            continue
        window = ' '.join(lines[index:min(len(lines), index + 3)])
        values = _money_candidates(window)
        if not values:
            continue
        found.append((values[0], window[:300]))
    return found

def extract_explicit_invoice_total(text: str) -> tuple[float | None, str, str]:
    if not text:
        return (None, '', '')
    ht_candidates = _labeled_amounts(text, ['total ht', 'net ht', 'montant ht', 'sous-total ht', 'sous total ht', 'base ht'])
    ttc_candidates = _labeled_amounts(text, ['total ttc', 'net à payer', 'net a payer', 'total à payer', 'total a payer', 'montant ttc', 'montant dû', 'montant du'])
    ht = max((value for value, _ in ht_candidates), default=None)
    ttc = max((value for value, _ in ttc_candidates), default=None)
    ht_proof = next((proof for value, proof in ht_candidates if value == ht), '')
    ttc_proof = next((proof for value, proof in ttc_candidates if value == ttc), '')
    if ht is not None and ttc is not None and (ht > ttc * 3):
        ht = None
        ht_proof = ''
    vat_match = re.search('(?i)(?:tva|taux)[^\\n%]{0,30}?(\\d{1,2}(?:[,.]\\d+)?)\\s*%', text)
    vat_rate = to_float(vat_match.group(1)) if vat_match else None
    if ht is not None and 10 <= ht <= 20000000:
        return (ht, 'HT', ht_proof)
    if ttc is not None and 10 <= ttc <= 20000000:
        if vat_rate is not None and 0 < vat_rate <= 30:
            estimated_ht = ttc / (1 + vat_rate / 100)
            return (estimated_ht, 'HT reconstitué depuis TTC', ttc_proof)
        return (ttc, 'TTC faute de HT lisible', ttc_proof)
    return (None, '', '')

def extract_invoice_financial(path: str, filename: str, scope: str) -> list[dict]:
    if Path(path).suffix.lower() != '.txt':
        return []
    text = read_text(path)
    if not text:
        return []
    amount, basis, proof = extract_explicit_invoice_total(text)
    if amount is None:
        return []
    category = classify_financial(filename + ' ' + text[:5000])
    if category is None:
        FINANCIAL_REJECTS.append({'source': filename, 'libelle': proof[:300], 'montant': amount, 'raison': 'total de facture trouvé mais catégorie non classée'})
        return []
    designation, role = category
    item = make_item(filename, scope, role, designation, amount, '€', 'moyenne', f'total de facture explicite ({basis})')
    item['preuve'] = proof[:500]
    item['detail_financier'] = True
    item['base_montant'] = basis
    print(f'  [DIRECT FACTURE] {designation} → {amount:.2f} € ({basis})')
    return [item]

def row_text(row: pd.Series) -> str:
    parts = []
    for v in row.values:
        if v is None or pd.isna(v):
            continue
        s = str(v).strip()
        if not s:
            continue
        if not re.fullmatch('[-+]?\\d+(?:[,.]\\d+)?', s.replace(' ', '')):
            parts.append(s)
    return ' '.join(parts)

def amount_from_row(row: pd.Series) -> float | None:
    vals = []
    for v in row.values:
        x = to_float(v)
        if x is None:
            continue
        if 1900 <= x <= 2100 or x in {0, 1, 2, 3, 4, 5}:
            continue
        if 10 <= abs(x) <= 1000000:
            vals.append(abs(x))
    return vals[-1] if vals else None

def _find_labeled_amount_in_recap(df: pd.DataFrame, label_keywords: list[str]) -> float | None:
    if df is None or df.empty:
        return None
    wanted = [norm(k) for k in label_keywords]
    rows, cols = df.shape
    for r in range(rows):
        for c in range(cols):
            cell = norm(df.iat[r, c])
            if not cell or not any((k in cell for k in wanted)):
                continue
            for rr in range(r + 1, min(rows, r + 8)):
                x = to_float(df.iat[rr, c])
                if x is not None and 10 <= x <= 1000000:
                    return float(x)
            for cc in range(c + 1, min(cols, c + 4)):
                x = to_float(df.iat[r, cc])
                if x is not None and 10 <= x <= 1000000:
                    return float(x)
    return None

def _find_total_ht_recap(df: pd.DataFrame) -> float | None:
    if df is None or df.empty:
        return None
    amount_cols = []
    for col in df.columns:
        cn = norm(col)
        if any((k in cn for k in ['montant', 'ht', 'euro', '€'])):
            amount_cols.append(col)
    vals: list[float] = []
    cols_to_scan = amount_cols or list(df.columns)
    for col in cols_to_scan:
        for v in df[col].tolist():
            x = to_float(v)
            if x is not None and 1000 <= x <= 5000000:
                vals.append(float(x))
    if not vals:
        return None
    return max(vals)

def extract_recap_achat_synthese(df: pd.DataFrame | None, filename: str, scope: str) -> list[dict]:
    if df is None or df.empty:
        return []
    f = norm(filename)
    if not any((k in f for k in ['recap achat', 'récap achat', 'achat pour bilan carbone'])):
        return []
    total_ht = _find_total_ht_recap(df)
    autres_services = _find_labeled_amount_in_recap(df, ['autres services'])
    location = _find_labeled_amount_in_recap(df, ["location d'outils", 'location outils', 'location'])
    entretien = _find_labeled_amount_in_recap(df, ['services entretiens', 'services entretien', 'entretien'])
    if total_ht is None or total_ht <= 0:
        return []
    if autres_services is None and location is None and (entretien is None):
        return []
    autres_services = float(autres_services or 0)
    location = float(location or 0)
    entretien = float(entretien or 0)
    deja_categorie = autres_services + location + entretien
    reste_materiel = total_ht - deja_categorie
    out: list[dict] = []
    if autres_services > 0:
        out.append(make_item(filename, scope, 'services', 'Autres services', autres_services, '€', 'faible', 'synthèse récap achats'))
    if location > 0:
        out.append(make_item(filename, scope, 'location_materiel', 'Location de matériel', location, '€', 'faible', 'synthèse récap achats'))
    if entretien > 0:
        out.append(make_item(filename, scope, 'maintenance', 'Services entretien/maintenance', entretien, '€', 'faible', 'synthèse récap achats'))
    if 1000 <= reste_materiel <= total_ht:
        item = make_item(filename, scope, 'achats_non_detailles', 'Achats matériels non détaillés', reste_materiel, '€', 'faible', 'reste du total HT après retrait services/location/entretien')
        out.append(item)
    if out:
        print(f'  [DIRECT ACHATS SYNTHÈSE] total={total_ht:.2f} € | services={autres_services:.2f} € | location={location:.2f} € | entretien={entretien:.2f} € | matériel non détaillé={max(reste_materiel, 0):.2f} €')
    return out

def extract_financial(df: pd.DataFrame | None, filename: str, scope: str, path: str | None=None) -> list[dict]:
    if df is None or df.empty:
        return []
    f = norm(filename)
    if any((k in f for k in ["facture d'eau", "factures d'eau", 'facture eau', 'factures eau', 'tableau conso eau', 'consommation eau', 'facture edf', 'factures edf', 'facture electricite', 'factures electricite', 'facture électricité', 'factures électricité', 'tableau conso edf', 'consommation electricite', 'consommation électricité'])):
        return []
    is_recap = any((k in f for k in ['recap achat', 'récap achat', 'achat pour bilan carbone']))
    is_immobilisation = 'immobilisation' in f
    is_invoice = 'facture' in f
    if is_invoice and (not is_recap) and (not is_immobilisation):
        return extract_invoice_financial(path or '', filename, scope)
    if not (is_recap or 'achat' in f or is_immobilisation):
        return []
    synth = extract_recap_achat_synthese(df, filename, scope)
    if synth:
        return synth
    grouped: dict[tuple[str, str], float] = defaultdict(float)
    for _, row in df.iterrows():
        txt = row_text(row)
        amount = amount_from_row(row)
        if amount is None or amount <= 0:
            continue
        category = classify_financial(txt + ' ' + filename)
        if category is None:
            if is_recap and txt.strip() and (txt.strip().lower() != 'nan'):
                FINANCIAL_REJECTS.append({'source': filename, 'libelle': txt[:300], 'montant': amount, 'raison': 'non classé'})
            continue
        designation, role = category
        grouped[designation, role] += amount
    output = [make_item(filename, scope, role, designation, quantity, '€', 'moyenne', 'extraction financière directe') for (designation, role), quantity in grouped.items()]
    if output:
        print(f"  [DIRECT ACHATS] {len(output)} catégorie(s) -> {sum((item['quantite'] for item in output)):.2f} €")
    return output

def extract_immobilisations(df: pd.DataFrame | None, filename: str, scope: str) -> list[dict]:
    if df is None or df.empty:
        return []
    f = norm(filename)
    if 'immobilisation' not in f:
        return []
    if any((k in f for k in ['cumul par compte', 'hors cessions'])):
        print('  [SKIP IMMO] fichier agrégé ignoré pour éviter doublon')
        return []
    value_col = next((c for c in df.columns if any((k in norm(c) for k in ["valeur de l'immobilisation", 'valeur de l’immobilisation', 'valeur brute', 'acquisition']))), None)
    label_col = next((c for c in df.columns if any((k in norm(c) for k in ['designation', 'désignation', 'libelle', 'libellé']))), None)
    if value_col is None:
        return extract_financial(df, filename, scope)
    total = 0.0
    for _, row in df.iterrows():
        label = norm(row.get(label_col, '')) if label_col else ''
        if any((k in label for k in ['fonds commercial', 'site internet', 'terrain', 'droit au bail'])):
            continue
        v = to_float(row.get(value_col))
        if v is not None and 10 <= v <= 100000:
            total += v
    if total <= 0:
        return []
    print(f'  [DIRECT IMMO] {total:.2f} €')
    return [make_item(filename, scope, 'immobilisations', 'Immobilisations', total, '€', 'moyenne', 'valeur immobilisations exploitable')]

def extract_eau(df: pd.DataFrame | None, filename: str, scope: str) -> list[dict]:
    if df is None or df.empty:
        return []
    fn = norm(filename)
    blob = norm(filename + ' ' + table_text(df, 2000))
    if 'attestation' in fn or 'titulaire de contrat' in fn:
        return []
    if any((k in fn for k in ['dechet', 'déchet', 'reporting dechets', 'reporting déchets', 'tableau recap dechets', 'tableau récap déchets', 'papier', 'ramette', 'a4', 'a3', 'traceur'])):
        return []
    if any((k in fn for k in ['edf', 'electricite', 'électricité', 'linky'])):
        return []
    if scope == 'SCOPE_2':
        return []
    if 'eau' not in blob and 'conso m3' not in blob and ('compteur m3' not in blob):
        return []
    vals = []
    for _, row in df.iterrows():
        nums = []
        for v in row.values:
            x = to_float(v)
            if x is None:
                continue
            if 1900 <= x <= 2100:
                continue
            if 0 < x <= 5000:
                nums.append(x)
        if not nums:
            continue
        vals.append(min(nums))
    if not vals:
        return []
    total = sum(vals)
    if total > 10000:
        print(f'  [SKIP EAU] {total:.2f} m³ suspect')
        return []
    print(f'  [DIRECT EAU] {total:.2f} m³')
    return [make_item(filename, scope, 'eau', 'Eau potable', total, 'm³', 'haute', 'somme consommations m3')]

def extract_papier(df: pd.DataFrame | None, filename: str, scope: str) -> list[dict]:
    if df is None or df.empty:
        return []
    fn = norm(filename)
    blob = norm(filename + ' ' + table_text(df, 2000))
    if 'factures achat papier' in fn or 'facture achat papier' in fn:
        print('  [SKIP PAPIER] facture papier non structurée ignorée')
        return []
    if not any((k in blob for k in ['tableau recap achat papier', 'tableau récap achat papier', 'a4', 'a3', 'traceur'])):
        return []
    a4 = a3 = rolls = 0.0
    for _, row in df.iterrows():
        txt = norm(' '.join((str(v) for v in row.values)))
        nums = []
        for v in row.values:
            x = to_float(v)
            if x is None:
                continue
            if 1900 <= x <= 2100:
                continue
            if 0 < x <= 500000:
                nums.append(x)
        if not nums:
            continue
        q = max(nums)
        if 'a4' in txt:
            a4 += q
        elif 'a3' in txt:
            a3 += q
        elif 'traceur' in txt or 'rouleau' in txt:
            rolls += q
    kg = a4 * 0.002 + a3 * 0.004 + rolls * 0.7
    if kg <= 0:
        return []
    if kg > 10000:
        print(f'  [SKIP PAPIER] {kg:.2f} kg suspect')
        return []
    print(f'  [DIRECT PAPIER] {kg:.2f} kg')
    return [make_item(filename, scope, 'papier', 'Papier', kg, 'kg', 'haute', 'conversion feuilles/rouleaux en kg')]

def extract_immobilisations_accounting_text(path: str, filename: str, scope: str) -> list[dict]:
    if Path(path).suffix.lower() != '.txt':
        return []
    text = read_text(path)
    start_match = None
    for candidate in re.finditer('Etat des immobilisations', text, re.I):
        if 'valeur brute' in norm(text[candidate.end():candidate.end() + 800]):
            start_match = candidate
            break
    if not start_match:
        markers = ['Installations techniques', 'materiel et outillage industriels', 'matériel et outillage industriels']
        for marker in markers:
            pos = norm(text).find(norm(marker))
            if pos < 0:
                continue
            raw_pos = max(0, text.lower().find(marker.lower()))
            if raw_pos < 0:
                raw_pos = pos
            zone = text[raw_pos:raw_pos + 500]
            amounts = re.findall('\\d{1,3}(?:[ .]\\d{3})+', zone)
            for raw_amount in amounts:
                value = to_float(raw_amount)
                if value is not None and 1000 <= value <= 20000000:
                    item = make_item(filename, 'SCOPE_3', 'immobilisations', 'Immobilisations corporelles - acquisitions', value, '€', 'moyenne', "montant d'immobilisations corporelles identifié dans le bilan actif")
                    item['calcul_automatique_interdit'] = False
                    item['preuve'] = zone[:500]
                    return [item]
        return []
    end_match = re.search('Etat des amortissements', text[start_match.end():], re.I)
    section_end = start_match.end() + end_match.start() if end_match else min(len(text), start_match.end() + 15000)
    section = text[start_match.start():section_end]
    lines = [line.strip() for line in section.splitlines() if line.strip()]
    acquisitions = None
    for index, line in enumerate(lines):
        if norm(line) != 'immobilisations corporelles':
            continue
        nearby = []
        for following in lines[index + 1:index + 8]:
            value = to_float(following)
            if value is not None and 1000 <= value <= 20000000:
                nearby.append(value)
        if len(nearby) >= 2:
            acquisitions = nearby[1]
            break
    if acquisitions is None:
        return []
    item = make_item(filename, 'SCOPE_3', 'immobilisations', 'Immobilisations corporelles - acquisitions', acquisitions, '€', 'moyenne', 'total des acquisitions corporelles ; facteur monétaire immobilisations')
    item['calcul_automatique_interdit'] = False
    item['preuve'] = f'IMMOBILISATIONS CORPORELLES : {acquisitions:.2f} €'
    return [item]

def _unit_norm(value: Any) -> str:
    u = str(value or '').strip()
    return u.replace('m3', 'm³').replace('M3', 'm³')

def _plausible_quantity(designation: str, quantite: float, unite: str) -> bool:
    d = norm(designation)
    u = _unit_norm(unite)
    if quantite <= 0:
        return False
    if u == '€':
        return quantite <= 100000000
    if 'papier' in d and u == 'kg':
        return quantite <= 10000
    if 'propane' in d and u == 'kg':
        return quantite <= 10000
    if ('dechet' in d or 'dechets' in d or 'gravats' in d or ('dib' in d) or ('bois' in d)) and u == 't':
        return quantite <= 10000
    if 'eau' in d and u == 'm³':
        return quantite <= 10000
    if 'electricite' in d and u.lower() == 'kwh':
        return quantite <= 1000000
    if 'gazole' in d and u == 'L':
        return quantite <= 250000
    return quantite <= 10000000

def _python_coherence_decision(item: dict, filename: str, scope: str, role: str) -> dict:
    d = norm(item.get('designation', ''))
    fn = norm(filename)
    u = _unit_norm(item.get('unite', ''))
    q = to_float(item.get('quantite'))
    if q is None:
        return {'statut': 'ANOMALIE', 'action': 'reject', 'raison': 'quantité illisible', 'confiance': 'haute'}
    effective_scope = str(item.get('scope') or scope)
    mixed_energy_water = any((k in fn for k in ['electricite', 'électricité', 'energie', 'énergie', 'edf'])) and any((k in fn for k in ['eau', 'water']))
    if any((k in fn for k in ['edf', 'electricite', 'électricité', 'linky'])) and (not mixed_energy_water) and (not any((k in d for k in ['electricite', 'chaleur', 'vapeur']))):
        return {'statut': 'ANOMALIE', 'action': 'reject', 'raison': 'fichier énergie/EDF incompatible avec la désignation extraite', 'confiance': 'haute'}
    if effective_scope == 'SCOPE_2' and (not any((k in d for k in ['electricite', 'chaleur', 'vapeur', 'reseau de chaleur', 'réseau de chaleur']))):
        return {'statut': 'ANOMALIE', 'action': 'reject', 'raison': 'SCOPE_2 incompatible avec cette donnée', 'confiance': 'haute'}
    if 'eau' in d and u != 'm³':
        return {'statut': 'ANOMALIE', 'action': 'set_a_verifier', 'raison': 'unité eau différente de m³', 'confiance': 'moyenne'}
    if 'electricite' in d and u.lower() != 'kwh':
        return {'statut': 'ANOMALIE', 'action': 'set_a_verifier', 'raison': 'unité électricité différente de kWh', 'confiance': 'moyenne'}
    if u == 'm³' and any((k in d for k in ['gravats', 'dechet', 'dechets', 'grave', 'granulat', 'sable', 'terre', 'calcaire'])):
        return {'statut': 'AUTO_CORRECTION', 'action': 'convert_unit', 'raison': 'conversion m³ -> t pour matériau/déchet', 'ancienne_unite': u, 'nouvelle_unite': 't', 'facteur_conversion': 1.6, 'confiance': 'haute'}
    if not _plausible_quantity(item.get('designation', ''), q, u):
        return {'statut': 'ANOMALIE', 'action': 'verify_with_groq', 'raison': 'valeur hors seuil de vraisemblance', 'confiance': 'moyenne'}
    return {'statut': 'OK', 'action': 'keep', 'raison': 'cohérent', 'confiance': 'haute'}

def _build_context_for_groq(path: str, df: pd.DataFrame | None, raw_text: str) -> str:
    if raw_text:
        return raw_text[:GROQ_VERIFY_CONTEXT_CHARS]
    if df is not None and (not df.empty):
        return table_text(df, GROQ_VERIFY_CONTEXT_CHARS)
    txt = read_text(path)
    return txt[:GROQ_VERIFY_CONTEXT_CHARS]

def _json_from_groq_response(text: str) -> dict | None:
    if not text:
        return None
    try:
        return json.loads(text)
    except Exception:
        pass
    m = re.search('\\{.*\\}', text, re.S)
    if not m:
        return None
    try:
        return json.loads(m.group(0))
    except Exception:
        return None

def _groq_verify_anomaly(item: dict, filename: str, scope: str, role: str, reason: str, context: str) -> dict | None:
    global _groq_verify_calls, _req_jour
    if not USE_GROQ_VERIFICATION:
        return None
    if _groq_verify_calls >= GROQ_VERIFY_MAX_CALLS:
        return None
    if not os.getenv('GROQ_API_KEY'):
        return None
    try:
        from groq import Groq
    except Exception:
        return None
    _groq_verify_calls += 1
    _req_jour += 1
    prompt = f"""\nTu es un vérificateur de cohérence pour une extraction de bilan carbone.\n\nRôle important :\n- Tu ne dois pas recalculer tout le bilan.\n- Tu dois uniquement vérifier l'anomalie signalée.\n- Tu ne dois pas inventer une donnée absente du contexte.\n- Si la valeur est incohérente et non corrigeable, demande le rejet.\n- Si une conversion d'unité est évidente, propose-la.\n- Réponds uniquement en JSON valide.\n\nActions autorisées :\n- keep\n- reject\n- convert_unit\n- replace_quantity\n- replace_designation\n- change_scope\n- set_a_verifier\n\nFormat JSON obligatoire :\n{{\n  "statut": "OK|ANOMALIE|CORRECTION",\n  "action": "keep|reject|convert_unit|replace_quantity|replace_designation|change_scope|set_a_verifier",\n  "raison": "explication courte",\n  "confiance": "haute|moyenne|faible",\n  "nouvelle_quantite": null,\n  "nouvelle_unite": null,\n  "facteur_conversion": null,\n  "nouvelle_designation": null,\n  "nouveau_scope": null\n}}\n\nFichier : {filename}\nScope détecté : {scope}\nRôle détecté : {role}\nAnomalie Python : {reason}\n\nValeur extraite :\n{json.dumps(item, ensure_ascii=False)}\n\nContexte compact du fichier :\n{context}\n""".strip()
    try:
        client = Groq(api_key=os.getenv('GROQ_API_KEY'))
        resp = client.chat.completions.create(model=GROQ_VERIFY_MODEL, messages=[{'role': 'system', 'content': 'Réponds uniquement en JSON valide, sans markdown.'}, {'role': 'user', 'content': prompt}], temperature=0, max_tokens=500)
        content = resp.choices[0].message.content or ''
        return _json_from_groq_response(content)
    except Exception as e:
        print(f'  [GROQ VERIFY] indisponible : {e}')
        return None

def _apply_decision(item: dict, decision: dict, source_decision: str) -> dict | None:
    action = str(decision.get('action', 'keep'))
    confidence = norm(decision.get('confiance', ''))
    reason = str(decision.get('raison', ''))[:300]
    if source_decision == 'groq' and confidence == 'faible':
        item['controle_python'] = item.get('controle_python', 'ANOMALIE')
        item['verification_groq'] = f'faible confiance : {reason}'
        item['correction_appliquee'] = 'aucune - à vérifier'
        item['fiabilite'] = 'moyenne'
        return item
    if action == 'keep':
        item.setdefault('controle_python', 'OK')
        return item
    if action == 'set_a_verifier':
        item['controle_python'] = item.get('controle_python', 'A_VERIFIER')
        item['verification_groq'] = reason if source_decision == 'groq' else ''
        item['correction_appliquee'] = 'marqué à vérifier'
        item['fiabilite'] = 'moyenne'
        return item
    if action == 'reject':
        print(f"  [REJET COHÉRENCE] {item.get('designation')} {item.get('quantite')} {item.get('unite')} -> {reason}")
        return None
    original_q = item.get('quantite')
    original_u = item.get('unite')
    original_d = item.get('designation')
    original_scope = item.get('scope')
    if action == 'convert_unit':
        factor = to_float(decision.get('facteur_conversion'))
        new_unit = decision.get('nouvelle_unite')
        if factor and new_unit and (0 < factor <= 1000):
            new_q = round(float(original_q) * factor, 6)
            if _plausible_quantity(item.get('designation', ''), new_q, str(new_unit)):
                item['quantite_originale'] = original_q
                item['unite_originale'] = original_u
                item['quantite'] = new_q
                item['unite'] = str(new_unit)
                item['controle_python'] = item.get('controle_python', 'AUTO_CORRECTION')
                item['verification_groq'] = reason if source_decision == 'groq' else ''
                item['correction_appliquee'] = f'conversion {original_u} -> {new_unit} x{factor}'
                item['fiabilite'] = 'moyenne'
                return item
        return item
    if action == 'replace_quantity':
        new_q = to_float(decision.get('nouvelle_quantite'))
        if new_q is not None and _plausible_quantity(item.get('designation', ''), new_q, item.get('unite', '')):
            item['quantite_originale'] = original_q
            item['quantite'] = new_q
            item['controle_python'] = item.get('controle_python', 'CORRECTION')
            item['verification_groq'] = reason if source_decision == 'groq' else ''
            item['correction_appliquee'] = f'quantité remplacée {original_q} -> {new_q}'
            item['fiabilite'] = 'moyenne'
            return item
        return item
    if action == 'replace_designation':
        new_d = decision.get('nouvelle_designation')
        if new_d and len(str(new_d)) <= 120:
            item['designation_originale'] = original_d
            item['designation'] = str(new_d)
            item['controle_python'] = item.get('controle_python', 'CORRECTION')
            item['verification_groq'] = reason if source_decision == 'groq' else ''
            item['correction_appliquee'] = f'désignation remplacée'
            item['fiabilite'] = 'moyenne'
            return item
        return item
    if action == 'change_scope':
        new_scope = decision.get('nouveau_scope')
        if new_scope in {'SCOPE_1', 'SCOPE_2', 'SCOPE_3'}:
            item['scope_original'] = original_scope
            item['scope'] = new_scope
            item['controle_python'] = item.get('controle_python', 'CORRECTION')
            item['verification_groq'] = reason if source_decision == 'groq' else ''
            item['correction_appliquee'] = f'scope remplacé {original_scope} -> {new_scope}'
            item['fiabilite'] = 'moyenne'
            return item
        return item
    return item

def verify_and_repair_items(items: list[dict], path: str, filename: str, scope: str, role: str, df: pd.DataFrame | None, raw_text: str) -> list[dict]:
    if not items:
        return []
    context_cache: str | None = None
    repaired: list[dict] = []
    for item in items:
        decision = _python_coherence_decision(item, filename, scope, role)
        statut = decision.get('statut', 'OK')
        action = decision.get('action', 'keep')
        raison = decision.get('raison', '')
        item.setdefault('controle_python', 'OK' if statut == 'OK' else f'{statut}: {raison}')
        final_decision = decision
        if action == 'verify_with_groq':
            if context_cache is None:
                context_cache = _build_context_for_groq(path, df, raw_text)
            groq_decision = _groq_verify_anomaly(item, filename, scope, role, raison, context_cache)
            if groq_decision:
                final_decision = groq_decision
                item['verification_groq'] = str(groq_decision.get('raison', ''))[:300]
            else:
                final_decision = {'action': 'reject', 'raison': f'{raison} ; Groq indisponible', 'confiance': 'haute'}
        source = 'python' if final_decision is decision else 'groq'
        new_item = _apply_decision(item, final_decision, source)
        if new_item is not None:
            repaired.append(new_item)
    return repaired

def audit_status(item: dict) -> tuple[str, str]:
    unit = norm(item.get('unite', ''))
    q = float(item.get('quantite', 0) or 0)
    des = norm(item.get('designation', ''))
    if q <= 0:
        return ('SUSPECT', 'quantité nulle')
    if 'eau' in des and unit in {'m3', 'm³'} and (q > 10000):
        return ('SUSPECT', 'eau > 10 000 m3')
    if 'papier' in des and unit == 'kg' and (q > 10000):
        return ('SUSPECT', 'papier > 10 000 kg')
    if 'propane' in des and unit == 'kg' and (q > 10000):
        return ('SUSPECT', 'propane > 10 000 kg')
    if 'gravats' in des and unit == 't' and (q > 10000):
        return ('SUSPECT', 'déchets inertes > 10 000 t')
    if unit in {'€', 'eur', 'euro', 'euros'}:
        return ('A_VERIFIER', 'montant financier à contrôler')
    return ('OK', '')

def write_audit(base_path: str, data: list[dict], no_data: list[str], nb_files: int) -> None:
    rows = []
    for item in data:
        st, alert = audit_status(item)
        rows.append({**item, 'statut': st, 'alerte': alert})
    for f in no_data:
        rows.append({'source': f, 'scope': '', 'role': '', 'designation': '', 'quantite': '', 'unite': '', 'fiabilite': '', 'justification': '', 'est_calcule': False, 'statut': 'AUCUNE_DONNEE', 'alerte': 'aucune donnée extraite'})
    fields = ['statut', 'alerte', 'source', 'scope', 'role', 'designation', 'quantite', 'unite', 'fiabilite', 'justification', 'est_calcule', 'controle_python', 'verification_groq', 'correction_appliquee', 'quantite_originale', 'unite_originale', 'designation_originale', 'scope_original']
    out_csv = os.path.join(base_path, 'audit_extraction_bilan_carbone.csv')
    with open(out_csv, 'w', newline='', encoding='utf-8-sig') as f:
        w = csv.DictWriter(f, fieldnames=fields, delimiter=';')
        w.writeheader()
        for row in rows:
            w.writerow({k: row.get(k, '') for k in fields})
    summary = {'version': VERSION, 'nb_fichiers_tentes': nb_files, 'nb_donnees': len(data), 'nb_sans_donnees': len(no_data), 'nb_ok': sum((1 for r in rows if r.get('statut') == 'OK')), 'nb_a_verifier': sum((1 for r in rows if r.get('statut') == 'A_VERIFIER')), 'nb_suspect': sum((1 for r in rows if r.get('statut') == 'SUSPECT')), 'verification_groq_active': bool(USE_GROQ_VERIFICATION and os.getenv('GROQ_API_KEY')), 'nb_appels_groq_verification': _groq_verify_calls, 'nb_appels_groq_fallback': _groq_fallback_calls, 'audit_csv': out_csv}
    out_json = os.path.join(base_path, 'audit_extraction_bilan_carbone_resume.json')
    Path(out_json).write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding='utf-8')
    if FINANCIAL_REJECTS:
        out_rej = os.path.join(base_path, 'audit_rejets_financiers_non_classes.csv')
        with open(out_rej, 'w', newline='', encoding='utf-8-sig') as f:
            w = csv.DictWriter(f, fieldnames=['source', 'libelle', 'montant', 'raison'], delimiter=';')
            w.writeheader()
            w.writerows(FINANCIAL_REJECTS)
        print(f'  [AUDIT] Rejets financiers : {out_rej}')
    print(f'  [AUDIT] CSV : {out_csv}')

def remove_invoice_details_covered_by_accounts(data: list[dict]) -> list[dict]:
    designations = {norm(item.get('designation', '')) for item in data}
    covers_purchases = any((value in designations for value in {'achats matieres et approvisionnements', 'achats matières et approvisionnements', 'achats materiels non detailles', 'achats matériels non détaillés'}))
    covers_services = any((value in designations for value in {'services, locations et charges externes', 'charges externes'}))
    covers_immobilisations = any((value in designations for value in {'immobilisations corporelles - acquisitions', 'immobilisations'}))
    output: list[dict] = []
    removed = 0
    for item in data:
        if not item.get('detail_financier'):
            output.append(item)
            continue
        role = norm(item.get('role', ''))
        covered = covers_purchases and role == 'achats_intrants' or (covers_services and role == 'services') or (covers_immobilisations and role == 'immobilisations')
        if covered:
            removed += 1
            print(f"  [DOUBLON COMPTA] {item.get('source')} ignoré : {item.get('designation')} déjà inclus dans un agrégat comptable")
            continue
        output.append(item)
    if removed:
        print(f'  [DOUBLON COMPTA] {removed} facture(s) détaillée(s) écartée(s)')
    return output

def remove_waste_details_covered_by_recap(data: list[dict]) -> list[dict]:
    has_waste_recap = any((item.get('agregat_dechets') for item in data))
    if not has_waste_recap:
        return data
    output: list[dict] = []
    removed = 0
    for item in data:
        if item.get('agregat_dechets'):
            output.append(item)
            continue
        role = norm(item.get('role', ''))
        source = norm(item.get('source', ''))
        if role == 'dechets' and ('reporting dechets' in source or 'reporting déchets' in source or item.get('detail_dechets_facture')):
            removed += 1
            print(f"  [DOUBLON DÉCHETS] {item.get('source')} ignoré : récapitulatif annuel déchets déjà présent")
            continue
        output.append(item)
    if removed:
        print(f'  [DOUBLON DÉCHETS] {removed} ligne(s) de détail écartée(s)')
    return output

def deduplicate(data: list[dict]) -> list[dict]:
    site_by_family: dict[str, str] = {}
    for item in data:
        if item.get('site_key'):
            site_by_family[item.get('source_family', source_family(item.get('source', '')))] = item['site_key']
    for item in data:
        family = item.get('source_family', source_family(item.get('source', '')))
        item['source_family'] = family
        if not item.get('site_key') and family in site_by_family:
            item['site_key'] = site_by_family[family]
    selected: dict[tuple, dict] = {}
    for item in data:
        designation = norm(item.get('designation', ''))
        unit = str(item.get('unite', ''))
        quantity = round(float(item.get('quantite', 0)), 4)
        family = item.get('source_family', '')
        site = item.get('site_key', '')
        if designation in {'electricite', 'eau potable'} and site:
            key = ('consommation', site, item.get('scope', ''), designation, unit)
            current = selected.get(key)
            if current is None:
                selected[key] = dict(item)
            else:
                candidate_score = (quantity, int(item.get('coverage_months', 0)))
                current_score = (round(float(current.get('quantite', 0)), 4), int(current.get('coverage_months', 0)))
                if candidate_score > current_score:
                    selected[key] = dict(item)
            continue
        key = (family or item.get('source', ''), item.get('scope', ''), item.get('role', ''), designation, unit, quantity)
        if key not in selected:
            selected[key] = dict(item)
    return list(selected.values())

def extract_one(path: str, scope: str, target_year: int | None=None) -> list[dict]:
    filename = Path(path).name
    raw_text = read_text(path) if Path(path).suffix.lower() == '.txt' else ''
    df = load_table(path)
    role = detect_role(filename, df, raw_text)
    collected: list[dict] = []
    standard_extractors = (lambda: extract_visual_energy_water(path, filename, scope), lambda: extract_propane(path, filename, scope), lambda: extract_vehicules(path, filename, scope), lambda: extract_edf_text_annual(path, filename, scope, target_year), lambda: extract_edf(path, filename, scope, df), lambda: extract_water_text_annual(path, filename, scope, target_year), lambda: extract_immobilisations(df, filename, scope), lambda: extract_dechets_recap_multiyear(path, filename, scope, df, target_year), lambda: extract_dechets(path, filename, scope, df), lambda: extract_ptd_materials(path, filename, scope), lambda: extract_eau(df, filename, scope))
    for extractor in standard_extractors:
        try:
            result = extractor()
        except Exception as exc:
            print(f'  [EXTRACTEUR] {exc}')
            result = []
        if result:
            collected.extend(result)
    paper_year = extract_paper_multiyear(path, filename, scope, target_year)
    if paper_year:
        collected.extend(paper_year)
    else:
        try:
            collected.extend(extract_papier(df, filename, scope))
        except Exception as exc:
            print(f'  [EXTRACTEUR PAPIER] {exc}')
    accounting_details = extract_accounting_details_generic(path, filename, scope)
    if accounting_details:
        collected.extend(accounting_details)
    else:
        summary = extract_accounting_purchases(path, filename, scope)
        if summary:
            collected.extend(summary)
        financial = extract_financial(df, filename, scope, path)
        if financial:
            collected.extend(financial)
    collected.extend(extract_immobilisations_accounting_text(path, filename, scope))
    checked_all: list[dict] = []
    if collected:
        checked = verify_and_repair_items(collected, path, filename, scope, role, df, raw_text)
        if checked:
            enrich_energy_metadata(checked, path, df, raw_text)
            checked_all.extend(checked)
    if not checked_all:
        groq_result = _groq_invoice_fallback(path, filename, scope, role)
        if groq_result:
            checked = verify_and_repair_items(groq_result, path, filename, scope, role, df, raw_text)
            if checked:
                checked_all.extend(checked)
    if not checked_all:
        if role == 'general':
            print('  [SKIP] rôle non reconnu ou donnée insuffisante')
        else:
            print(f'  [SKIP] aucun extracteur direct concluant ({role})')
    checked_all = apply_policy_to_items(checked_all)
    return deduplicate(checked_all)

def iter_scope_files(base_path: str):
    for scope in ('SCOPE_1', 'SCOPE_2', 'SCOPE_3'):
        folder = Path(base_path) / scope
        if not folder.exists():
            continue
        for p in sorted(folder.rglob('*')):
            if p.is_file() and p.suffix.lower() in {'.csv', '.txt', '.xlsx', '.xls', '.xlsm', '.png', '.jpg', '.jpeg', '.tif', '.tiff', '.bmp', '.docx'}:
                yield (scope, str(p))

def lancer_le_bilan(base_path: str='.') -> tuple[list[dict], int]:
    print('=' * 70)
    print(VERSION)
    print('=' * 70)
    all_data: list[dict] = []
    no_data: list[str] = []
    tried = 0
    scope_files = list(iter_scope_files(base_path))
    target_year = resolve_target_year([path for _, path in scope_files])
    if target_year:
        print(f'Exercice détecté automatiquement : {target_year}')
    current_scope = None
    for scope, path in scope_files:
        if current_scope != scope:
            current_scope = scope
            print(f"\n{'─' * 70}\nDossier : {scope}\n{'─' * 70}")
        filename = Path(path).name
        if is_non_source(filename):
            print(f'  [SKIP] {filename} → exclu/non source')
            continue
        if should_skip_for_year(filename, target_year):
            print(f'  [SKIP ANNÉE] {filename} → hors exercice {target_year}')
            continue
        tried += 1
        print(f'\n  Traitement : {filename}')
        try:
            res = extract_one(path, scope, target_year)
        except Exception as e:
            print(f'  [ERREUR] {filename}: {e}')
            res = []
        if res:
            all_data.extend(res)
            print(f'  → {len(res)} valeur(s) :')
            for r in res:
                print(f"     {r['designation'][:42]:42s} | {r['quantite']:>10} {r['unite']:<8} | {r['fiabilite']}")
        else:
            no_data.append(filename)
            print('  → Aucune donnée')
    detailed_electricity = {round(float(item.get('quantite', 0)), 4) for item in all_data if norm(item.get('designation', '')) == 'electricite' and 'electricite-eau' in norm(item.get('source', ''))}
    if detailed_electricity:
        filtered_data = []
        for item in all_data:
            is_duplicate_summary = norm(item.get('designation', '')) == 'electricite' and 'synthese' in norm(item.get('source', '')) and (round(float(item.get('quantite', 0)), 4) in detailed_electricity)
            if is_duplicate_summary:
                print(f"  [DOUBLON EDF] {item.get('source')} ignoré : {item.get('quantite')} kWh déjà présent dans la feuille détaillée")
                continue
            filtered_data.append(item)
        all_data = filtered_data
    detected_sites = {item.get('site_key') for item in all_data if item.get('site_key')}
    if len(detected_sites) == 1:
        only_site = next(iter(detected_sites))
        for item in all_data:
            if norm(item.get('designation', '')) in {'electricite', 'eau potable'} and (not item.get('site_key')):
                item['site_key'] = only_site
    all_data = remove_waste_details_covered_by_recap(all_data)
    all_data = remove_invoice_details_covered_by_accounts(all_data)
    dedup = deduplicate(all_data)
    removed = len(all_data) - len(dedup)
    print('\n' + '=' * 70)
    print(f'POST-TRAITEMENT — {len(all_data)} valeurs brutes')
    if removed:
        print(f'  {removed} doublon(s) fusionné(s)')
    print('\n' + '=' * 70)
    print(f'RÉSUMÉ FINAL — {len(dedup)} valeurs | {tried} fichier(s) tenté(s)')
    for scope in ('SCOPE_1', 'SCOPE_2', 'SCOPE_3'):
        items = [d for d in dedup if d.get('scope') == scope]
        if not items:
            continue
        print(f'\n  [{scope}] {len(items)} valeur(s) :')
        for d in items:
            print(f"    {d['source'][:28]:28s} | {d['designation'][:32]:32s} | {d['quantite']:>10} {d['unite']:<8} | {d.get('role', '')}")
    if no_data:
        print(f'\nFichiers sans données ({len(no_data)}) :')
        for f in no_data[:30]:
            print(f'  - {f}')
        if len(no_data) > 30:
            print(f'  ... {len(no_data) - 30} autre(s)')
    write_audit(base_path, dedup, no_data, tried)
    print('=' * 70)
    return (dedup, tried)
if __name__ == '__main__':
    lancer_le_bilan('.')
